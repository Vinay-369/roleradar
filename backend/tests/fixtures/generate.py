"""
Generates real PDF/DOCX fixtures for resume-parsing tests, so the tests
exercise actual PyMuPDF/python-docx extraction rather than fake strings.
"""
import fitz
from docx import Document

GOOD_RESUME_TEXT = """Ananya Rao
ananya.rao@example.com | 9876543210 | github.com/ananyarao | linkedin.com/in/ananyarao

Summary
Final-year Computer Science student focused on backend development.

Skills
Python, FastAPI, MongoDB, Docker, REST APIs, Git

Experience
Backend Intern, StartupX
Built REST APIs used by 3 internal teams.
Reduced API response time by optimizing database queries.

Projects
Resume Intelligence Tool
Built a Flask app that parses resumes and scores ATS compatibility.

Education
B.Tech Computer Science, XYZ University, 2026

Certifications
AWS Cloud Practitioner

Achievements
Won first place in college hackathon 2025
"""


def make_good_pdf(path: str):
    doc = fitz.open()
    page = doc.new_page()
    # Single column: write the whole text as one block down the left side.
    page.insert_text((50, 50), GOOD_RESUME_TEXT, fontsize=10)
    doc.save(path)
    doc.close()


def make_two_column_pdf(path: str):
    doc = fitz.open()
    page = doc.new_page()
    left_col = "Skills\nPython\nFastAPI\nMongoDB\nDocker\n" * 3
    right_col = "Experience\nBackend Intern\nBuilt REST APIs\nOptimized queries\n" * 3
    # Two side-by-side columns at the same y-range -> should trip the
    # multi-column heuristic.
    page.insert_text((40, 50), left_col, fontsize=10)
    page.insert_text((320, 50), right_col, fontsize=10)
    doc.save(path)
    doc.close()


def make_scanned_like_pdf(path: str):
    """Near-empty text layer, simulating a scanned/image-only resume."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "img", fontsize=8)
    doc.save(path)
    doc.close()


def make_good_docx(path: str):
    document = Document()
    for line in GOOD_RESUME_TEXT.split("\n"):
        document.add_paragraph(line)
    document.save(path)


def make_docx_with_table(path: str):
    document = Document()
    document.add_paragraph("Ananya Rao")
    document.add_paragraph("ananya.rao@example.com 9876543210")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, MongoDB")
    document.add_paragraph("Experience")
    document.add_paragraph("Education")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Company"
    document.save(path)


if __name__ == "__main__":
    import os
    here = os.path.dirname(__file__)
    make_good_pdf(os.path.join(here, "good_resume.pdf"))
    make_two_column_pdf(os.path.join(here, "two_column_resume.pdf"))
    make_scanned_like_pdf(os.path.join(here, "scanned_like_resume.pdf"))
    make_good_docx(os.path.join(here, "good_resume.docx"))
    make_docx_with_table(os.path.join(here, "docx_with_table.docx"))
    print("fixtures generated")
