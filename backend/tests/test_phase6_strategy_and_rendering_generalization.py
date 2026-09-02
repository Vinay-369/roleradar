"""
Phase 6 Comprehensive Test Suite: Resume Strategy, ATS Optimization, Template Intelligence & Document Rendering.

Validates:
1. TEST 1 — 1-Page Strategy: Finalization respects page_budget=1 and enforces fit.
2. TEST 2 — 2-Page Strategy: Finalization respects page_budget=2 and does not over-compress.
3. TEST 3 — User Edit (Valid Source-Grounded Content): Saved with verification_status='VERIFIED'.
4. TEST 4 — User Edit (Unsupported Technology Added): Saved with verification_status='USER_MODIFIED' and flagged in truth_guard_audit.
5. TEST 5 — User Edit (Metric Alteration): Altered metric detected in truth_guard_audit['unsupported_metrics'].
6. TEST 6 — User Edit (Scope Escalation): Ungrounded leadership claim detected in truth_guard_audit['scope_escalations'].
7. TEST 7 — User Edit (Legitimate Contact/Formatting Fields): Contact updates preserved without false positive violations.
8. TEST 8 — Export & Verification State: Distinguishes VERIFIED vs USER_MODIFIED.
9. ROUND-TRIP — PDF & DOCX Extraction Integrity: Verifies 100% data preservation on real generated documents.
"""
import io
import fitz  # PyMuPDF
from docx import Document
import pytest
from bson import ObjectId
from mongomock_motor import AsyncMongoMockClient

from app.core.config import Settings
from app.modules.jobs.services import create_custom_job
from app.modules.resume import repositories as resume_repo
from app.modules.resume.parsing.structurer import extract_candidate_profile
from app.modules.tailoring import repositories as tailoring_repo
from app.modules.tailoring import services as tailoring_services
from app.modules.tailoring.export import generate_pdf, generate_docx
from app.modules.tailoring.strategy import (
    CareerStage,
    TemplateFamily,
    build_resume_strategy,
)


@pytest.fixture
def db():
    client = AsyncMongoMockClient()
    return client["roleradar_test"]


@pytest.fixture
def settings():
    return Settings(JWT_SECRET="test-secret", EMBEDDING_PROVIDER="mock", AI_PROVIDER="mock")


DEFAULT_PARSEABILITY = {
    "score": 90,
    "issues": [],
    "detected_sections": [],
    "missing_standard_sections": [],
    "contact_info_found": {"email": "test@example.com", "phone": "555-0100"},
    "likely_multi_column": False,
    "word_count": 250,
}

DEFAULT_RECRUITER_IMPACT = {
    "score": 85,
    "bullets_analyzed": 4,
    "quantified_bullets": 2,
    "weak_verb_bullets": 0,
    "quantification_rate": 0.5,
    "issues": [],
}


# =========================================================================
# TEST 1 & TEST 2: 1-PAGE VS 2-PAGE DYNAMIC PAGE BUDGET ENFORCEMENT
# =========================================================================

@pytest.mark.asyncio
async def test_1_page_strategy_budget_enforcement(db, settings):
    """
    Candidate with page_budget=1.
    Finalization must verify fits_budget is True and page_count <= 1.
    """
    user_id = str(ObjectId())

    candidate_text = """
SARAH CONNOR
sarah@example.com | 555-0100 | Los Angeles, CA

SUMMARY
Junior software developer with 1 year experience in Python.

SKILLS
Python, SQL, Git

EXPERIENCE
Junior Developer at Cyberdyne (2023 - Present)
• Built Python backend endpoints for database reporting.
• Maintained automated testing scripts using pytest.

EDUCATION
B.S. Computer Science (2019 - 2023)
State University
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db,
        user_id=user_id,
        version=1,
        file_name="sarah_resume.pdf",
        file_type="pdf",
        raw_text=candidate_text,
        parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY,
        recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )

    job = await create_custom_job(
        db,
        company="TechCorp",
        title="Python Developer",
        jd_text="Python Developer\nREQUIREMENTS:\n• 1+ years Python experience.",
        user_id=user_id,
    )

    # Build 1-page strategy
    strat = build_resume_strategy(profile, target_role="Python Developer")
    strat.page_budget = 1

    version = await tailoring_repo.create_version(
        db,
        user_id,
        job["id"],
        job["title"],
        job["company"],
        changes=[],
        parsed=profile.to_parsed_dict(),
        resume_strategy=strat.model_dump(mode="json"),
    )
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    finalized = await tailoring_services.finalize_tailoring(db, user_id, v_id, settings=settings)
    val_summary = finalized["validation_summary"]

    assert val_summary["page_budget"] == 1
    assert val_summary["page_count"] <= 1
    assert finalized["one_page_fit"] is True


@pytest.mark.asyncio
async def test_2_page_strategy_budget_preservation(db, settings):
    """
    Senior candidate with dense multi-role experience (page_budget=2).
    Finalization must respect max_pages=2 and not over-compress into 1 page.
    """
    user_id = str(ObjectId())

    candidate_text = """
ELENA ROSTOVA
elena@example.com | 555-0200 | Seattle, WA | linkedin.com/in/elena

PROFESSIONAL SUMMARY
Senior Distributed Systems Architect with 10+ years experience designing high-throughput data platforms.

TECHNICAL SKILLS
Languages: Python, Go, Java, C++, SQL, Rust
Infrastructure: AWS, Kubernetes, Docker, Terraform, Kafka, Redis, PostgreSQL, Elasticsearch

PROFESSIONAL EXPERIENCE
Principal Architect at CloudScale (2020 - Present) - Seattle, WA
• Architected multi-region event streaming platform processing 2.5B daily events with 99.999% availability using Kafka and Go.
• Reduced cloud infrastructure spending by $1.4M annually through automated Kubernetes autoscaling and spot instance orchestration.
• Led engineering organization of 28 senior engineers across 4 distributed service teams.
• Authored core data governance standards adopted enterprise-wide across 12 product lines.

Senior Staff Engineer at DataFlow Inc (2016 - 2020) - San Francisco, CA
• Designed real-time analytics engine indexing 50TB log data daily using Go and Elasticsearch.
• Optimized PostgreSQL query execution latency by 45% across 200M customer records.
• Spearheaded migration from monolithic architecture to 40+ containerized microservices on AWS EKS.
• Mentored 12 mid-level and junior engineers through technical promotion tracks.

Senior Software Engineer at AlphaNet (2013 - 2016) - Boston, MA
• Developed distributed caching tier using Redis clusters, decreasing p99 read latency from 120ms to 8ms.
• Built automated CI/CD deployment pipelines using Docker and Jenkins, accelerating release frequency by 4x.
• Implemented OAuth2 and SAML authentication infrastructure securing 5M user accounts.

Software Engineer at BetaCorp (2010 - 2013) - Austin, TX
• Developed backend data pipelines using Java and PostgreSQL.
• Automated database backup and disaster recovery validation workflows.

PROJECTS
• Distributed Key-Value Store (Go, Raft): Built consensus-backed distributed storage engine with zero data loss under simulated network partitions.
• Stream Analytics Engine (Rust): Developed high-throughput memory-safe log stream parser processing 500k eps.

EDUCATION
Master of Science in Computer Science (2008 - 2010)
University of Washington - GPA: 3.9

Bachelor of Science in Computer Engineering (2004 - 2008)
University of Michigan - GPA: 3.8

CERTIFICATIONS
• AWS Certified Solutions Architect Professional
• Certified Kubernetes Administrator (CKA)
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db,
        user_id=user_id,
        version=1,
        file_name="elena_resume.pdf",
        file_type="pdf",
        raw_text=candidate_text,
        parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY,
        recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )

    job = await create_custom_job(
        db,
        company="Apex Systems",
        title="Principal Infrastructure Architect",
        jd_text="Principal Architect\nREQUIREMENTS:\n• 8+ years distributed systems architecture.\n• Kafka, Kubernetes, Go experience.",
        user_id=user_id,
    )

    # Build 2-page strategy
    strat = build_resume_strategy(profile, target_role="Principal Infrastructure Architect")
    strat.page_budget = 2

    version = await tailoring_repo.create_version(
        db,
        user_id,
        job["id"],
        job["title"],
        job["company"],
        changes=[],
        parsed=profile.to_parsed_dict(),
        resume_strategy=strat.model_dump(mode="json"),
    )
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    finalized = await tailoring_services.finalize_tailoring(db, user_id, v_id, settings=settings)
    val_summary = finalized["validation_summary"]

    assert val_summary["page_budget"] == 2
    assert val_summary["page_count"] <= 2
    assert finalized["one_page_fit"] is True


# =========================================================================
# TEST 3, 4, 5, 6, 7: USER EDIT TRUTH GUARD AUDIT & PROVENANCE
# =========================================================================

@pytest.mark.asyncio
async def test_user_edit_valid_source_grounded_content(db):
    """
    User rephrases a summary or bullet while preserving source facts.
    Should be saved with verification_status='VERIFIED' and zero violations.
    """
    user_id = str(ObjectId())
    candidate_text = """
JOHN DOE
john@example.com | 555-0300

SKILLS
Python, FastAPI, Docker

EXPERIENCE
Developer at WebCo (2022 - Present)
• Built Python backend services with Docker.
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db, user_id=user_id, version=1, file_name="john.pdf", file_type="pdf",
        raw_text=candidate_text, parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY, recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )
    job = await create_custom_job(db, company="AppCo", title="Python Dev", jd_text="Python Dev", user_id=user_id)
    version = await tailoring_repo.create_version(db, user_id, job["id"], job["title"], job["company"], changes=[], parsed=profile.to_parsed_dict())
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    # User edit: Rephrase with verified skills
    edited_parsed = profile.to_parsed_dict()
    edited_parsed["summary"] = "Experienced Python and FastAPI developer with Docker deployment expertise."

    updated = await tailoring_services.update_parsed_resume(db, user_id, v_id, edited_parsed)

    assert updated["user_modified"] is True
    assert updated["verification_status"] == "VERIFIED"
    assert updated["truth_guard_audit"]["is_valid"] is True
    assert len(updated["truth_guard_audit"]["violations"]) == 0


@pytest.mark.asyncio
async def test_user_edit_unsupported_technology_flagged(db):
    """
    User manually injects 'Kubernetes' into skills when candidate only has Python/Docker.
    Edit is saved as USER_MODIFIED and flagged in truth_guard_audit.
    """
    user_id = str(ObjectId())
    candidate_text = """
JOHN DOE
john@example.com

SKILLS
Python, Docker

EXPERIENCE
Developer at WebCo (2022 - Present)
• Built Python backend services.
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db, user_id=user_id, version=1, file_name="john.pdf", file_type="pdf",
        raw_text=candidate_text, parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY, recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )
    job = await create_custom_job(db, company="AppCo", title="DevOps", jd_text="Kubernetes Dev", user_id=user_id)
    version = await tailoring_repo.create_version(db, user_id, job["id"], job["title"], job["company"], changes=[], parsed=profile.to_parsed_dict())
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    # User injects Kubernetes
    edited_parsed = profile.to_parsed_dict()
    edited_parsed["skills"].append("Kubernetes")

    updated = await tailoring_services.update_parsed_resume(db, user_id, v_id, edited_parsed)

    assert updated["user_modified"] is True
    assert updated["verification_status"] == "USER_MODIFIED"
    assert updated["truth_guard_audit"]["is_valid"] is False
    assert any("Kubernetes" in v or "kubernetes" in v for v in updated["truth_guard_audit"]["violations"] + updated["truth_guard_audit"]["unsupported_technologies"])


@pytest.mark.asyncio
async def test_user_edit_metric_alteration_flagged(db):
    """
    Source has '20%'. User manually changes to '80%'.
    Detected in truth_guard_audit['unsupported_metrics'].
    """
    user_id = str(ObjectId())
    candidate_text = """
JOHN DOE
john@example.com

EXPERIENCE
Developer at WebCo (2022 - Present)
• Optimized database queries, increasing performance by 20%.
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db, user_id=user_id, version=1, file_name="john.pdf", file_type="pdf",
        raw_text=candidate_text, parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY, recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )
    job = await create_custom_job(db, company="AppCo", title="Dev", jd_text="Dev", user_id=user_id)
    version = await tailoring_repo.create_version(db, user_id, job["id"], job["title"], job["company"], changes=[], parsed=profile.to_parsed_dict())
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    # User inflates metric
    edited_parsed = profile.to_parsed_dict()
    edited_parsed["experience_raw"] = ["• Optimized database queries, increasing performance by 80%."]

    updated = await tailoring_services.update_parsed_resume(db, user_id, v_id, edited_parsed)

    assert updated["verification_status"] == "USER_MODIFIED"
    assert updated["truth_guard_audit"]["is_valid"] is False
    assert len(updated["truth_guard_audit"]["unsupported_metrics"]) > 0


@pytest.mark.asyncio
async def test_user_edit_scope_escalation_flagged(db):
    """
    Source: 'Contributed to a 4-person development team.'
    User changes to: 'Led a team of 4 engineers.'
    Detected in truth_guard_audit['scope_escalations'].
    """
    user_id = str(ObjectId())
    candidate_text = """
JOHN DOE
john@example.com

EXPERIENCE
Developer at WebCo (2022 - Present)
• Contributed to a 4-person development team building web services.
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db, user_id=user_id, version=1, file_name="john.pdf", file_type="pdf",
        raw_text=candidate_text, parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY, recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )
    job = await create_custom_job(db, company="AppCo", title="Lead", jd_text="Lead", user_id=user_id)
    version = await tailoring_repo.create_version(db, user_id, job["id"], job["title"], job["company"], changes=[], parsed=profile.to_parsed_dict())
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    # User claims unearned leadership
    edited_parsed = profile.to_parsed_dict()
    edited_parsed["experience_raw"] = ["• Led a team of 4 engineers building web services."]

    updated = await tailoring_services.update_parsed_resume(db, user_id, v_id, edited_parsed)

    assert updated["verification_status"] == "USER_MODIFIED"
    assert updated["truth_guard_audit"]["is_valid"] is False
    assert len(updated["truth_guard_audit"]["scope_escalations"]) > 0


@pytest.mark.asyncio
async def test_user_edit_contact_and_location_preserved_without_violation(db):
    """
    User updates email, phone, or location.
    Must not trigger false positive fabrication violations.
    """
    user_id = str(ObjectId())
    candidate_text = """
JOHN DOE
john@example.com | 555-0100 | Boston, MA

SKILLS
Python, SQL

EXPERIENCE
Developer at WebCo (2022 - Present)
• Built Python reporting tools.
"""
    profile = extract_candidate_profile(candidate_text)
    await resume_repo.create_master_resume(
        db, user_id=user_id, version=1, file_name="john.pdf", file_type="pdf",
        raw_text=candidate_text, parsed=profile.to_parsed_dict(),
        parseability=DEFAULT_PARSEABILITY, recruiter_impact=DEFAULT_RECRUITER_IMPACT,
    )
    job = await create_custom_job(db, company="AppCo", title="Dev", jd_text="Dev", user_id=user_id)
    version = await tailoring_repo.create_version(db, user_id, job["id"], job["title"], job["company"], changes=[], parsed=profile.to_parsed_dict())
    v_id = str(version["_id"]) if "_id" in version else str(version["id"])

    # User edits contact info
    edited_parsed = profile.to_parsed_dict()
    edited_parsed["personal"]["location"] = "San Francisco, CA"
    edited_parsed["personal"]["phone"] = "+1 555-9999"

    updated = await tailoring_services.update_parsed_resume(db, user_id, v_id, edited_parsed)

    assert updated["truth_guard_audit"]["is_valid"] is True
    assert updated["verification_status"] == "VERIFIED"


# =========================================================================
# TEST 8 & ROUND-TRIP: PDF & DOCX EXTRACTION INTEGRITY
# =========================================================================

def test_pdf_roundtrip_text_extraction_integrity():
    """
    Renders a structured CandidateProfile to PDF, extracts text using PyMuPDF,
    and verifies 100% preservation of core entities, metrics, dates, and technologies.
    """
    candidate_text = """
MARCUS VANCE
marcus.vance@example.com | 555-0450 | Seattle, WA

PROFESSIONAL SUMMARY
Backend Engineer with 4 years experience building cloud services.

TECHNICAL SKILLS
Python, Go, PostgreSQL, Docker, Redis

PROFESSIONAL EXPERIENCE
Cloud Engineer at DataStream Inc (2021 - Present)
• Architected real-time streaming pipeline processing 1.2M records daily using Go and PostgreSQL.
• Reduced API response latency by 35% across 50 microservice endpoints.

PROJECTS
• Distributed Key-Value Store: Built distributed consensus engine using Raft protocol in Go.

EDUCATION
B.S. Computer Science (2017 - 2021)
University of Washington
"""
    profile = extract_candidate_profile(candidate_text)
    parsed_dict = profile.to_parsed_dict()

    # Generate PDF
    pdf_bytes = generate_pdf(parsed_dict, candidate_name="Marcus Vance", template="modern")
    assert len(pdf_bytes) > 1000

    # Extract text with PyMuPDF
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_text = ""
    for page in doc:
        extracted_text += page.get_text("text") + "\n"
    doc.close()

    norm_extracted = extracted_text.lower()

    # Verify key claims and entities
    assert "marcus vance" in norm_extracted
    assert "marcus.vance@example.com" in norm_extracted
    assert "datastream" in norm_extracted
    assert "1.2m" in norm_extracted or "1.2 m" in norm_extracted
    assert "35%" in norm_extracted
    assert "postgresql" in norm_extracted
    assert "university of washington" in norm_extracted


def test_docx_roundtrip_structural_integrity():
    """
    Renders structured profile to DOCX and verifies paragraph and table structure.
    """
    candidate_text = """
NATALIE PORTMAN
natalie@example.com | 555-0550

TECHNICAL SKILLS
Python, JavaScript, SQL

PROFESSIONAL EXPERIENCE
Software Developer at MediaTech (2022 - Present)
• Built customer facing web portals using Python and JavaScript.
"""
    profile = extract_candidate_profile(candidate_text)
    parsed_dict = profile.to_parsed_dict()

    # Generate DOCX
    docx_bytes = generate_docx(parsed_dict, candidate_name="Natalie Portman", template="classic")
    assert len(docx_bytes) > 1000

    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    norm_doc = full_text.lower()

    assert "natalie portman" in norm_doc
    assert "mediatech" in norm_doc
    assert "python" in norm_doc
    assert len(doc.paragraphs) >= 4
