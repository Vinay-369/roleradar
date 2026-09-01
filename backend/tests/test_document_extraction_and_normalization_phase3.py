"""
Tests for Phase 3: Document Extraction and Normalization.
Validates layout extraction, multi-column reading order, Unicode normalization,
line unwrapping, DOCX extraction, table extraction, and factual metric preservation.
"""
import io
import fitz  # PyMuPDF
import pytest
from docx import Document

from app.modules.resume.parsing.document import (
    DocumentBlock,
    DocumentPage,
    NormalizedDocument,
)
from app.modules.resume.parsing.normalization import (
    normalize_unicode_artifacts,
    unwrap_paragraph_lines,
    detect_and_order_columns,
    suppress_repeated_headers_footers,
)
from app.modules.resume.parsing.text_extraction import (
    extract_pdf,
    extract_docx,
    extract_text_and_layout,
)


def _create_synthetic_pdf(pages_content: list[list[tuple[float, float, str]]], page_size=(612, 792)) -> bytes:
    """Helper to generate in-memory synthetic PDF with explicit (x, y, text) coordinates."""
    doc = fitz.open()
    for items in pages_content:
        page = doc.new_page(width=page_size[0], height=page_size[1])
        for x, y, text in items:
            page.insert_text((x, y), text, fontsize=10)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _create_synthetic_docx(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Helper to generate in-memory synthetic DOCX."""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if tables:
        for t_data in tables:
            t = doc.add_table(rows=len(t_data), cols=len(t_data[0]) if t_data else 0)
            for r_idx, row in enumerate(t_data):
                for c_idx, cell_value in enumerate(row):
                    t.cell(r_idx, c_idx).text = cell_value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_unicode_and_bullet_normalization():
    raw_text = (
        "Name: Alex Taylor\u00a0\u00a0\u00a0\u200b\n"
        "Summary: Senior Engineer with \u2018cloud\u2019 and \u201cmicroservices\u201d experience.\n"
        "▪ Spearheaded micro-\nservices migration reducing latency by 45%.\n"
        "► Architected distributed queue handling 10,000+ msgs/sec.\n"
        "◆ Deployed Kubernetes clusters with 99.99% uptime.\n"
        "✓ Cut release cycles from 3 hours \u2013 10 minutes.\n"
    )

    cleaned = normalize_unicode_artifacts(raw_text)

    # 1. Check Unicode artifacts stripped
    assert "\u00a0" not in cleaned
    assert "\u200b" not in cleaned
    assert "‘" not in cleaned and "’" not in cleaned
    assert "“" not in cleaned and "”" not in cleaned

    # 2. Check hyphenated word unwrapping
    assert "microservices migration" in cleaned

    # 3. Check bullet standardization
    lines = cleaned.splitlines()
    bullet_lines = [l for l in lines if l.startswith("•")]
    assert len(bullet_lines) == 4

    # 4. Check strict factual metric preservation
    assert "45%" in cleaned
    assert "10,000+" in cleaned
    assert "99.99%" in cleaned
    assert "3 hours" in cleaned and "10 minutes" in cleaned


def test_line_unwrapping_preserves_separate_items_and_headings():
    wrapped_lines = [
        "WORK EXPERIENCE",
        "• Engineered resilient message delivery platform on AWS SQS and",
        "  Redis streams processing 50,000 requests per minute with",
        "  zero message loss.",
        "• Led cross-functional team of 5 engineers.",
        "Developer Tools:",
        "• Built automated CI/CD pipeline cutting build times by 40%.",
    ]

    unwrapped = unwrap_paragraph_lines(wrapped_lines)

    assert len(unwrapped) == 5
    assert unwrapped[0] == "WORK EXPERIENCE"
    assert "Engineered resilient message delivery platform on AWS SQS and Redis streams processing 50,000 requests per minute with zero message loss." in unwrapped[1]
    assert unwrapped[2] == "• Led cross-functional team of 5 engineers."
    assert "Developer Tools:" in unwrapped[3]
    assert "• Built automated CI/CD pipeline cutting build times by 40%." in unwrapped[4]


def test_multi_column_reading_order_inference():
    # Synthetic two-column page layout
    # Header: top full width (x=50, y=50)
    # Left Column: x=50, y=100 (Work Experience)
    # Right Column: x=350, y=100 (Technical Skills)
    b_header = DocumentBlock(id="b0", page=0, text="ALEX TAYLOR\nBangalore | alex@example.com", bbox=(50, 50, 550, 80))
    b_left_1 = DocumentBlock(id="b1", page=0, text="WORK EXPERIENCE\nSenior Software Engineer at TechCorp", bbox=(50, 100, 280, 150))
    b_left_2 = DocumentBlock(id="b2", page=0, text="• Scaled backend to 100k QPS", bbox=(50, 160, 280, 200))
    b_right_1 = DocumentBlock(id="b3", page=0, text="TECHNICAL SKILLS\nPython, Go, Docker, Kubernetes", bbox=(350, 100, 550, 150))
    b_right_2 = DocumentBlock(id="b4", page=0, text="EDUCATION\nB.Tech Computer Science", bbox=(350, 160, 550, 200))

    # Scrambled input order (e.g. standard left-to-right interleaving from naive extraction)
    scrambled = [b_header, b_left_1, b_right_1, b_left_2, b_right_2]

    ordered = detect_and_order_columns(scrambled, page_width=612.0)

    # Column ordering should place: Header -> Left Col (b1, b2) -> Right Col (b3, b4)
    ordered_ids = [b.id for b in ordered]
    assert ordered_ids == ["b0", "b1", "b2", "b3", "b4"]


def test_repeated_headers_and_footers_suppression():
    # Two-page document where "Alex Taylor — Resume" and "Page 1 of 2" / "Page 2 of 2" appear at margins
    p0_blocks = [
        DocumentBlock(id="p0_h", page=0, text="Alex Taylor — Resume", bbox=(50, 20, 550, 40)),
        DocumentBlock(id="p0_b", page=0, text="Experience content on page 1", bbox=(50, 100, 550, 300)),
        DocumentBlock(id="p0_f", page=0, text="Page 1 of 2", bbox=(250, 760, 350, 780)),
    ]
    p1_blocks = [
        DocumentBlock(id="p1_h", page=1, text="Alex Taylor — Resume", bbox=(50, 20, 550, 40)),
        DocumentBlock(id="p1_b", page=1, text="Projects content on page 2", bbox=(50, 100, 550, 300)),
        DocumentBlock(id="p1_f", page=1, text="Page 2 of 2", bbox=(250, 760, 350, 780)),
    ]

    pages = [
        DocumentPage(page_number=0, width=612.0, height=792.0, blocks=p0_blocks),
        DocumentPage(page_number=1, width=612.0, height=792.0, blocks=p1_blocks),
    ]

    suppress_repeated_headers_footers(pages)

    # Top header should be marked as repeated
    assert p0_blocks[0].is_repeated_header_or_footer is True
    assert p1_blocks[0].is_repeated_header_or_footer is True

    # Main content should NOT be marked
    assert p0_blocks[1].is_repeated_header_or_footer is False
    assert p1_blocks[1].is_repeated_header_or_footer is False

    # Page numbers should be marked
    assert p0_blocks[2].is_repeated_header_or_footer is True
    assert p1_blocks[2].is_repeated_header_or_footer is True


def test_synthetic_pdf_extraction_to_normalized_document():
    pdf_bytes = _create_synthetic_pdf([
        [
            (50, 50, "JORDAN SMITH"),
            (50, 70, "jordan@example.com | +1 555-0199 | San Francisco, CA"),
            (50, 120, "EXPERIENCE"),
            (50, 140, "Staff Infrastructure Engineer at CloudCore (2021 - Present)"),
            (50, 160, "• Architected globally distributed multi-region database serving 500k QPS with 99.999% availability."),
            (50, 180, "• Reduced AWS compute costs by $450,000 annually through ARM64 instance migration."),
        ]
    ])

    result = extract_text_and_layout(pdf_bytes, "jordan_resume.pdf")

    assert result["file_type"] == "pdf"
    assert result["is_scanned"] is False
    assert "JORDAN SMITH" in result["text"]
    assert "500k QPS" in result["text"]
    assert "$450,000" in result["text"]

    norm_doc = result["document"]
    assert isinstance(norm_doc, NormalizedDocument)
    assert len(norm_doc.pages) == 1
    assert len(norm_doc.blocks) >= 1

    # Check that backward compatible legacy blocks are populated
    assert len(result["blocks"]) >= 1
    assert all("page" in b and "x0" in b and "text" in b for b in result["blocks"])


def test_synthetic_docx_extraction_with_tables_to_normalized_document():
    paragraphs = [
        "ELENA ROSTOVA",
        "elena@example.com | Seattle, WA",
        "SUMMARY",
        "Principal Machine Learning Systems Engineer with 10+ years experience.",
        "EXPERIENCE",
        "Principal Engineer at DeepAI (2020 - Present)",
        "• Engineered distributed LLM training framework scaling to 2,048 H100 GPUs.",
    ]
    tables = [
        [
            ["Skill Category", "Technologies"],
            ["Languages", "Python, C++, CUDA, Rust"],
            ["Frameworks", "PyTorch, TensorRT, Triton, vLLM"],
        ]
    ]

    docx_bytes = _create_synthetic_docx(paragraphs, tables)

    result = extract_text_and_layout(docx_bytes, "elena_resume.docx")

    assert result["file_type"] == "docx"
    assert result["has_tables"] is True
    assert "ELENA ROSTOVA" in result["text"]
    assert "2,048 H100 GPUs" in result["text"]
    assert "TensorRT" in result["text"]

    norm_doc = result["document"]
    assert isinstance(norm_doc, NormalizedDocument)
    assert norm_doc.has_tables is True

    # Check table blocks exist in document
    table_blocks = [b for b in norm_doc.blocks if b.is_table_cell]
    assert len(table_blocks) >= 4
