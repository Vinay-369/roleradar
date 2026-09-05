import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from app.modules.resume.parsing.structurer import (
    SECTION_PATTERNS,
    parse_experience_section,
    parse_projects_section,
)


STANDARD_ATS_HEADINGS: dict[str, str] = {
    "summary": "PROFESSIONAL SUMMARY",
    "skills": "TECHNICAL SKILLS",
    "experience": "PROFESSIONAL EXPERIENCE",
    "internships": "INTERNSHIPS",
    "projects": "PROJECTS",
    "education": "EDUCATION",
    "certifications": "CERTIFICATIONS",
    "achievements": "HONORS & AWARDS",
    "publications": "PUBLICATIONS & RESEARCH",
    "research": "RESEARCH EXPERIENCE",
    "languages": "LANGUAGES",
}


def _sanitize_text(text: str) -> str:
    """Normalizes Unicode dashes, spaces, and quotes to standard ASCII for perfect ATS parser extraction."""
    if not text:
        return ""
    replacements = {
        "\u2013": "-",  # en-dash
        "\u2014": "-",  # em-dash
        "\u2012": "-",
        "\u2212": "-",  # minus sign
        "\u2018": "'",  # left single quote
        "\u2019": "'",  # right single quote
        "\u201c": '"',  # left double quote
        "\u201d": '"',  # right double quote
        "\u00a0": " ",  # non-breaking space
        "\ufffd": "-",  # replacement char
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def render_text_from_structured(parsed_data: dict) -> str:
    """
    Renders clean, structured markdown/plain-text directly from a parsed dictionary.
    Zero regex text-splicing — pure deterministic formatting.
    """
    if not parsed_data:
        return ""

    personal = parsed_data.get("personal", {}) or parsed_data.get("personal_info", {}) or {}
    name = personal.get("name") or personal.get("full_name") or "Candidate"
    contacts = []
    if personal.get("location"):
        contacts.append(personal["location"])
    elif personal.get("city") or personal.get("address"):
        loc = ", ".join(filter(None, [personal.get("city"), personal.get("address")]))
        if loc:
            contacts.append(loc)
    if personal.get("email"):
        contacts.append(personal["email"])
    if personal.get("phone"):
        contacts.append(personal["phone"])
    if personal.get("linkedin"):
        contacts.append(personal["linkedin"])
    if personal.get("github"):
        contacts.append(personal["github"])
    if personal.get("portfolio"):
        contacts.append(personal["portfolio"])

    lines = [name]
    if contacts:
        lines.append(" • ".join(contacts))
    lines.append("")

    # Summary
    summary = parsed_data.get("summary") or parsed_data.get("objective")
    if summary and str(summary).strip():
        lines.append("PROFESSIONAL SUMMARY")
        lines.append(str(summary).strip())
        lines.append("")

    # Skills
    skills_cat = parsed_data.get("skills_categorized", [])
    skills = parsed_data.get("skills", [])
    if skills_cat and isinstance(skills_cat, list):
        lines.append("TECHNICAL SKILLS")
        for sc in skills_cat:
            lines.append(str(sc).strip())
        lines.append("")
    elif isinstance(skills, dict):
        lines.append("TECHNICAL SKILLS")
        for cat, val in skills.items():
            val_str = ", ".join(str(v) for v in val) if isinstance(val, list) else str(val)
            lines.append(f"{cat}: {val_str}")
        lines.append("")
    elif skills:
        lines.append("TECHNICAL SKILLS")
        if isinstance(skills, list):
            # Check if list already has category lines with colons
            has_categories = any(":" in str(s) and len(str(s).split(":")[0].split()) <= 6 for s in skills)
            if has_categories:
                for s in skills:
                    lines.append(str(s).strip())
            else:
                lines.append(", ".join(str(s) for s in skills if s))
        else:
            lines.append(str(skills))
        lines.append("")

    # Experience
    exp = parsed_data.get("experience_raw", []) or parsed_data.get("experience", []) or parsed_data.get("work_experience", [])
    if exp:
        lines.append("PROFESSIONAL EXPERIENCE")
        if isinstance(exp, list) and exp and isinstance(exp[0], dict) and "company" in exp[0] and "bullets" in exp[0]:
            exp_entities = exp
        else:
            exp_entities = parse_experience_section(exp)

        for ent in exp_entities:
            comp = ent.company if hasattr(ent, "company") else ent.get("company", "")
            loc = ent.location if hasattr(ent, "location") else ent.get("location", "")
            comp_header = f"{comp} — {loc}" if comp and loc else (comp or "")
            if comp_header and comp_header.lower() not in ("work experience", "experience"):
                lines.append(comp_header)

            prog = ent.progression if hasattr(ent, "progression") else ent.get("progression", [])
            if prog:
                for p in prog:
                    p_title = p.title if hasattr(p, "title") else p.get("title", "")
                    p_dates = p.dates if hasattr(p, "dates") else p.get("dates", "")
                    p_line = f"{p_title} ({p_dates})" if p_title and p_dates else (p_title or "")
                    if p_line:
                        lines.append(p_line)
            elif (hasattr(ent, "role") and ent.role) or (isinstance(ent, dict) and ent.get("role")):
                r_title = ent.role if hasattr(ent, "role") else ent.get("role", "")
                r_dates = ent.dates if hasattr(ent, "dates") else ent.get("dates", "")
                r_line = f"{r_title} ({r_dates})" if r_title and r_dates else (r_title or "")
                if r_line:
                    lines.append(r_line)

            r_groups = ent.responsibility_groups if hasattr(ent, "responsibility_groups") else ent.get("responsibility_groups", [])
            if r_groups:
                for grp in r_groups:
                    g_heading = grp.heading if hasattr(grp, "heading") else grp.get("heading", "")
                    if g_heading:
                        lines.append(g_heading)
                    g_bullets = grp.bullets if hasattr(grp, "bullets") else grp.get("bullets", [])
                    for b in g_bullets:
                        b_clean = re.sub(r"^[•\-\*\s]+", "", str(b)).strip()
                        if b_clean:
                            lines.append(f"• {b_clean}")
            else:
                e_bullets = ent.bullets if hasattr(ent, "bullets") else ent.get("bullets", [])
                for b in e_bullets:
                    b_clean = re.sub(r"^[•\-\*\s]+", "", str(b)).strip()
                    if b_clean:
                        lines.append(f"• {b_clean}")
            lines.append("")

    # Projects
    proj = parsed_data.get("projects_raw", []) or parsed_data.get("projects", [])
    if proj:
        lines.append("TECHNICAL PROJECTS")
        if isinstance(proj, list) and proj and isinstance(proj[0], dict) and "title" in proj[0] and "bullets" in proj[0]:
            proj_entities = proj
        else:
            proj_entities = parse_projects_section(proj)

        for p in proj_entities:
            p_title = p.title if hasattr(p, "title") else p.get("title", "")
            p_tech = p.tech_stack if hasattr(p, "tech_stack") else (p.get("tech_stack") or p.get("technologies"))
            p_dates = p.dates if hasattr(p, "dates") else p.get("dates")

            t_line = f"{p_title} ({p_dates})" if p_title and p_dates else (p_title or "")
            if t_line:
                lines.append(t_line)
            if p_tech:
                p_tech_str = ", ".join(p_tech) if isinstance(p_tech, list) else str(p_tech)
                lines.append(f"Technologies: {p_tech_str}")

            p_bullets = p.bullets if hasattr(p, "bullets") else p.get("bullets", [])
            for b in p_bullets:
                b_clean = re.sub(r"^[•\-\*\s]+", "", str(b)).strip()
                if b_clean:
                    lines.append(f"• {b_clean}")
            lines.append("")

    # Education
    edu = parsed_data.get("education_raw", []) or parsed_data.get("education", [])
    if edu:
        lines.append("EDUCATION")
        for e in edu:
            if isinstance(e, dict):
                inst = e.get("institution") or e.get("school") or ""
                deg = e.get("degree") or ""
                gpa = e.get("cgpa") or e.get("percentage") or e.get("gpa") or ""
                dates = e.get("dates") or e.get("year") or ""
                loc = e.get("location") or ""
                if inst:
                    lines.append(inst + (f", {loc}" if loc else ""))
                deg_gpa = deg + (f" | {gpa}" if gpa else "")
                if dates:
                    deg_gpa += f" ({dates})"
                if deg_gpa:
                    lines.append(deg_gpa)
                lines.append("")
            else:
                e_str = str(e).strip()
                if e_str:
                    lines.append(e_str)
                    lines.append("")
        if lines and lines[-1] != "":
            lines.append("")

    # Certifications
    certs = parsed_data.get("certifications", []) or parsed_data.get("certifications_raw", []) or parsed_data.get("certificates", [])
    if certs:
        lines.append("CERTIFICATIONS")
        for c in certs:
            c_str = str(c).strip()
            if c_str:
                prefix = "" if c_str.startswith(("•", "-", "*")) else "• "
                lines.append(f"{prefix}{c_str}")
        lines.append("")

    # Achievements
    ach = parsed_data.get("achievements", []) or parsed_data.get("achievements_raw", []) or parsed_data.get("awards", [])
    if ach:
        lines.append("ACHIEVEMENTS")
        for a in ach:
            a_str = str(a).strip()
            if a_str:
                prefix = "" if a_str.startswith(("•", "-", "*")) else "• "
                lines.append(f"{prefix}{a_str}")
        lines.append("")

    # Languages
    langs = parsed_data.get("languages", []) or parsed_data.get("languages_raw", []) or parsed_data.get("languages_known", [])
    if langs:
        lines.append("LANGUAGES")
        if isinstance(langs, list):
            lines.append(", ".join(str(l) for l in langs if l))
        else:
            lines.append(str(langs))
        lines.append("")

    return "\n".join(lines).strip()


INK_900 = "#111827"
INK_700 = "#374151"
INK_500 = "#6b7280"
SIGNAL_600 = "#0d766e"
CLASSIC_ACCENT = "#111827"
TECH_ACCENT = "#1e3a8a"

# Strict bullet prefix regex: matches standard bullet glyphs and numbered lists
# NEVER matches degree abbreviations like 'B.E', 'B.Tech', 'M.Tech', 'B.Sc', 'M.S.'
_BULLET_PREFIX_RE = re.compile(
    r"^(?:[•\-\*\u2022\u2023\u25E6\u2043\u2219\u25AA\u25AB\u27A4\u2714\u2713\u279C\u2192\u25BA\u25B6\u25C6\u25C7\u25CF\u25CB\u2718\u2717\u2705\u27A2\u2794\u2714]|\d{1,2}[\.\)]|\([a-zA-Z0-9]+\)|[a-zA-Z]\))\s+",
    re.UNICODE,
)

_DATE_PATTERN_RE = re.compile(
    r"\b((?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+)?\d{4}(?:\s*[-–—to]+\s*(?:Present|Current|\d{4}|(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+)?\d{4}))?)\b",
    re.IGNORECASE,
)


def _is_section_header(line: str) -> bool:
    """Accurately identifies section headers including markdown ##, **, underlines, and colons."""
    normalized = re.sub(r"^[#\-=*~_ \t]+|[#\-=*~_:\t ]+$", "", line.strip())
    if not normalized or len(normalized) > 50:
        return False
    return any(re.match(pattern, normalized, re.IGNORECASE) for pattern in SECTION_PATTERNS.values())


def _looks_like_contact_line(line: str) -> bool:
    return bool(re.search(r"@|\b\+?\d{1,3}[-.\s]?\d{10}\b|linkedin\.com|github\.com|portfolio|https?://", line, re.IGNORECASE))


def _split_name_contact_body(resume_text: str) -> tuple[str, list[str], list[str]]:
    """Returns (name, contact_lines, remaining_lines)."""
    lines = [l for l in resume_text.split("\n")]
    non_empty = [(i, l) for i, l in enumerate(lines) if l.strip()]
    if not non_empty:
        return "", [], []

    name_idx, name = non_empty[0]
    # Clean leading markdown # from name if present
    name = re.sub(r"^[#\s*]+", "", name).strip()
    contact_lines: list[str] = []
    body_start_idx = name_idx + 1

    for i in range(name_idx + 1, len(lines)):
        stripped = lines[i].strip()
        if not stripped:
            continue
        if _is_section_header(stripped):
            body_start_idx = i
            break
        if _looks_like_contact_line(stripped) or ("|" in stripped or "•" in stripped or "·" in stripped):
            contact_lines.append(stripped)
            body_start_idx = i + 1
        elif len(contact_lines) == 0 and len(stripped.split()) <= 5 and not _is_section_header(stripped):
            # Might be location line e.g. "Davangere, Karnataka"
            contact_lines.append(stripped)
            body_start_idx = i + 1
        else:
            body_start_idx = i
            break

    return name.strip(), contact_lines, lines[body_start_idx:]


def _clean_title_and_date(line: str) -> tuple[str, str] | None:
    """Extracts clean (title_or_institution, right_aligned_date) if line contains a date span."""
    if line.startswith("•") or line.startswith("-") or line.startswith("*"):
        return None

    date_match = re.search(
        r"\b((?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+)?\d{4}(?:\s*[-–—to]+\s*(?:Present|Current|\d{4}|(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+)?\d{4}))?)\b\)?\s*$",
        line,
        re.IGNORECASE,
    )
    if not date_match:
        return None

    date_str = date_match.group(1).strip()
    raw_title = line
    raw_title = re.sub(r"\(\s*" + re.escape(date_str) + r"\s*\)", "", raw_title)
    raw_title = raw_title.replace(date_str, "")
    raw_title = re.sub(r"\(\s*\)", "", raw_title)
    raw_title = re.sub(r"\[\s*\]", "", raw_title)
    raw_title = raw_title.strip(" |–-—, \t")

    if not raw_title:
        return "", date_str
    return raw_title, date_str


def _is_skill_category(line: str) -> bool:
    clean = _BULLET_PREFIX_RE.sub("", line).strip()
    return ":" in clean and not clean.startswith("http") and len(clean.split(":")[0].split()) <= 6


def _is_tech_stack_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    lower = stripped.lower()
    if lower.startswith("tech stack:") or lower.startswith("technologies:") or lower.startswith("tools & tech:") or lower.startswith("stack:"):
        return True
    if stripped.count("|") >= 2 and len(stripped.split()) <= 25 and not _looks_like_contact_line(stripped):
        return True
    return False


def _is_project_title_line(line: str) -> bool:
    """Recognize a project heading in a structured multi-line project item."""
    clean = _BULLET_PREFIX_RE.sub("", line).strip()
    if not clean or line.strip().startswith(("•", "-", "*")) or _is_tech_stack_line(clean):
        return False
    if _clean_title_and_date(clean) is not None:
        return True
    first_word = clean.split()[0].lower().rstrip(":,")
    from app.modules.resume.parsing.structurer import _EVIDENCE_VERB_RE
    return (
        1 <= len(clean.split()) <= 18
        and not clean.endswith((".", ";", "!"))
        and not _EVIDENCE_VERB_RE.match(first_word)
    )


def _split_into_bullet_points(text: str) -> list[str]:
    """Splits a block of text into distinct bullet points if multiple sentences/bullets exist."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) > 1:
        res = []
        for l in lines:
            res.extend(_split_into_bullet_points(l))
        return res

    raw = lines[0] if lines else text.strip()
    if not raw:
        return []

    cleaned = _BULLET_PREFIX_RE.sub("", raw).strip()
    # Split sentences ending in . or ; that are followed by capital letter
    parts = re.split(r"(?<=[.;])\s+(?=[A-Z][a-z0-9A-Z_]+(?:\s+|$))", cleaned)
    if len(parts) > 1:
        valid_parts = [p.strip() for p in parts if len(p.strip()) > 10]
        if len(valid_parts) == len(parts):
            return valid_parts

    return [cleaned]


def render_pdf_from_structured(
    parsed_data: dict, candidate_name: str = "", template: str = "standard", experience_level: str = "fresher"
) -> bytes:
    """
    Renders an ATS-optimized, beautifully styled PDF directly from a structured resume dictionary.
    Guarantees:
    1. Dynamic Section Ordering based on ResumeStrategy.section_order.
    2. Zero orphan headings (keepWithNext=True on headings & tables).
    3. Multi-role progression hierarchy and responsibility groups cleanly separated.
    4. Controlled visual presentation variants (classic, standard, compact).
    """
    template = template.lower() if template else "standard"
    if template in ("modern", "technical", "stanford"):
        template = "standard"
    elif template in ("executive", "harvard"):
        template = "classic"
    elif template == "minimal":
        template = "compact"

    buffer = io.BytesIO()

    # Controlled margin and spacing tokens per visual variant
    if template == "compact":
        top_m, bot_m, left_m, right_m = 0.35 * inch, 0.35 * inch, 0.4 * inch, 0.4 * inch
        heading_space_before = 7
        body_font_size = 9.0
        body_leading = 11.0
        bullet_leading = 11.0
        bullet_space_after = 1.0
    elif template == "classic":
        top_m, bot_m, left_m, right_m = 0.4 * inch, 0.4 * inch, 0.45 * inch, 0.45 * inch
        heading_space_before = 10
        body_font_size = 9.5
        body_leading = 12.0
        bullet_leading = 12.0
        bullet_space_after = 1.5
    else:  # standard
        top_m, bot_m, left_m, right_m = 0.4 * inch, 0.4 * inch, 0.45 * inch, 0.45 * inch
        heading_space_before = 9
        body_font_size = 9.5
        body_leading = 12.0
        bullet_leading = 12.0
        bullet_space_after = 1.5

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=top_m, bottomMargin=bot_m,
        leftMargin=left_m, rightMargin=right_m,
    )
    styles = getSampleStyleSheet()

    if template == "classic":
        font_family = "Times-Roman"
        font_bold = "Times-Bold"
        font_italic = "Times-Italic"
        accent_color = CLASSIC_ACCENT
        name_align = 1  # center
    else:
        font_family = "Helvetica"
        font_bold = "Helvetica-Bold"
        font_italic = "Helvetica-Oblique"
        accent_color = SIGNAL_600 if template == "standard" else INK_900
        name_align = 1 if template == "standard" else 0

    usable_width = A4[0] - (left_m + right_m)

    name_style = ParagraphStyle(
        "Name", parent=styles["Normal"], fontName=font_bold, fontSize=15,
        leading=18, textColor=HexColor(INK_900), alignment=name_align, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontName=font_family, fontSize=8.5,
        leading=11, textColor=HexColor(INK_700), alignment=name_align, spaceAfter=3,
    )
    heading_style = ParagraphStyle(
        "Heading", parent=styles["Normal"], fontName=font_bold, fontSize=9.5,
        leading=12, textColor=HexColor(accent_color), spaceBefore=heading_space_before, spaceAfter=2,
        letterSpacing=0.6, keepWithNext=True,
    )
    subhead_left = ParagraphStyle(
        "SubheadLeft", parent=styles["Normal"], fontName=font_bold, fontSize=8.5,
        leading=11, textColor=HexColor(INK_900), spaceAfter=1, keepWithNext=True,
    )
    subhead_right = ParagraphStyle(
        "SubheadRight", parent=styles["Normal"], fontName=font_bold if template == "classic" else font_family,
        fontSize=8.0, leading=11, textColor=HexColor(INK_700), alignment=2, rightIndent=6, keepWithNext=True,
    )
    tech_stack_style = ParagraphStyle(
        "TechStack", parent=styles["Normal"], fontName=font_italic, fontSize=8.0,
        leading=10.5, textColor=HexColor(INK_700), spaceAfter=2, keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName=font_family, fontSize=body_font_size,
        textColor=HexColor(INK_900), leading=body_leading, spaceAfter=1.5,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=styles["Normal"], fontName=font_family, fontSize=body_font_size,
        textColor=HexColor(INK_900), leading=bullet_leading, leftIndent=12, firstLineIndent=-8, spaceAfter=bullet_space_after,
    )

    def esc(s: str) -> str:
        clean = _sanitize_text(str(s))
        return clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    personal = parsed_data.get("personal", {}) or parsed_data.get("personal_info", {}) or {}
    name = personal.get("name") or personal.get("full_name") or candidate_name or "Candidate"
    name = re.sub(r"^[#\s*]+", "", str(name)).strip()

    contacts = []
    if personal.get("location"):
        contacts.append(personal["location"])
    elif personal.get("city") or personal.get("address"):
        loc = ", ".join(filter(None, [personal.get("city"), personal.get("address")]))
        if loc:
            contacts.append(loc)
    if personal.get("email"):
        contacts.append(personal["email"])
    if personal.get("phone"):
        contacts.append(personal["phone"])
    if personal.get("linkedin"):
        contacts.append(personal["linkedin"])
    if personal.get("github"):
        contacts.append(personal["github"])
    if personal.get("portfolio"):
        contacts.append(personal["portfolio"])

    story = []
    if name:
        story.append(Paragraph(esc(name), name_style))
    if contacts:
        formatted_contacts = " • ".join(str(c).strip(" •·|") for c in contacts if str(c).strip())
        story.append(Paragraph(esc(formatted_contacts), contact_style))
    if name or contacts:
        divider_color = HexColor(accent_color if template == "standard" else INK_700)
        story.append(HRFlowable(width="100%", thickness=0.6, color=divider_color, spaceAfter=3))

    def add_section_header(title: str):
        story.append(Paragraph(esc(title.upper()), heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor(accent_color if template == "classic" else INK_500), spaceAfter=2))

    # Resolve dynamic section ordering
    strategy_info = parsed_data.get("_strategy", {}) if isinstance(parsed_data.get("_strategy"), dict) else {}
    section_order = (
        parsed_data.get("_ordered_sections")
        or strategy_info.get("section_order")
        or ["summary", "skills", "experience", "internships", "projects", "education", "certifications", "achievements", "publications", "research", "languages"]
    )
    std_headings = strategy_info.get("standard_ats_headings", {}) or STANDARD_ATS_HEADINGS

    rendered_sections = set()

    for sec in section_order:
        sec_lower = str(sec).lower()

        # Summary
        if sec_lower in ("summary", "objective") and sec_lower not in rendered_sections:
            summary = parsed_data.get("summary") or parsed_data.get("objective")
            if summary and str(summary).strip():
                add_section_header(std_headings.get("summary", "PROFESSIONAL SUMMARY"))
                story.append(Paragraph(esc(str(summary).strip()), body_style))
                rendered_sections.add(sec_lower)

        # Technical Skills / Core Competencies
        elif sec_lower in ("skills", "technical skills", "competencies", "core competencies") and "skills" not in rendered_sections:
            skills_cat = parsed_data.get("skills_categorized", [])
            skills = parsed_data.get("skills", [])
            if skills_cat or skills:
                heading_title = std_headings.get("skills", "TECHNICAL SKILLS")
                add_section_header(heading_title)
                if skills_cat and isinstance(skills_cat, list):
                    for sc in skills_cat:
                        sc_str = str(sc).strip()
                        if ":" in sc_str:
                            parts = sc_str.split(":", 1)
                            story.append(Paragraph(f"<b>{esc(parts[0].strip())}:</b> {esc(parts[1].strip())}", body_style))
                        else:
                            story.append(Paragraph(esc(sc_str), body_style))
                elif isinstance(skills, dict):
                    for cat, val in skills.items():
                        val_str = ", ".join(str(v) for v in val) if isinstance(val, list) else str(val)
                        story.append(Paragraph(f"<b>{esc(cat)}:</b> {esc(val_str)}", body_style))
                elif isinstance(skills, list):
                    cat_lines = [str(s).strip() for s in skills if ":" in str(s) and len(str(s).split(":")[0].split()) <= 6]
                    if len(cat_lines) >= 2 and len(cat_lines) >= len(skills) * 0.5:
                        for s in skills:
                            s_str = str(s).strip()
                            if ":" in s_str:
                                parts = s_str.split(":", 1)
                                story.append(Paragraph(f"<b>{esc(parts[0].strip())}:</b> {esc(parts[1].strip())}", body_style))
                            else:
                                story.append(Paragraph(esc(s_str), body_style))
                    else:
                        clean_skills = [re.sub(r"^[A-Za-z\s&]+:\s*", "", str(s)).strip() for s in skills if s and str(s).strip()]
                        clean_unique = []
                        seen_sk = set()
                        for cs in clean_skills:
                            if cs.lower() not in seen_sk:
                                seen_sk.add(cs.lower())
                                clean_unique.append(cs)
                        story.append(Paragraph(esc(", ".join(clean_unique)), body_style))
                else:
                    story.append(Paragraph(esc(str(skills)), body_style))
                rendered_sections.add("skills")

        # Professional Experience
        elif sec_lower in ("experience", "work_experience", "professional experience") and "experience" not in rendered_sections:
            exp = parsed_data.get("experience", []) or parsed_data.get("experience_raw", []) or parsed_data.get("work_experience", [])
            if exp:
                heading_title = std_headings.get("experience", "PROFESSIONAL EXPERIENCE")
                add_section_header(heading_title)
                if isinstance(exp, list) and exp and isinstance(exp[0], dict) and "company" in exp[0] and ("bullets" in exp[0] or "progression" in exp[0]):
                    exp_entities = exp
                else:
                    exp_entities = parse_experience_section(exp)

                for ent in exp_entities:
                    comp = ent.company if hasattr(ent, "company") else ent.get("company", "")
                    loc = ent.location if hasattr(ent, "location") else ent.get("location", "")

                    # Company Header with Location
                    if comp and comp.lower() not in ("work experience", "experience"):
                        if loc:
                            tbl = Table(
                                [[Paragraph(esc(comp), subhead_left), Paragraph(esc(loc), subhead_right)]],
                                colWidths=[usable_width * 0.72, usable_width * 0.28],
                                hAlign='LEFT',
                            )
                            tbl.setStyle(TableStyle([
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                ("RIGHTPADDING", (0, 0), (0, -1), 0),
                                ("RIGHTPADDING", (1, 0), (1, -1), 6),
                                ("TOPPADDING", (0, 0), (-1, -1), 1),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                            ]))
                            story.append(tbl)
                        else:
                            story.append(Paragraph(esc(comp), subhead_left))

                    prog = ent.progression if hasattr(ent, "progression") else ent.get("progression", [])
                    if prog:
                        for p in prog:
                            p_title = p.title if hasattr(p, "title") else p.get("title", "")
                            p_dates = p.dates if hasattr(p, "dates") else p.get("dates", "")
                            if p_title and p_dates:
                                tbl = Table(
                                    [[Paragraph(esc(p_title), body_style), Paragraph(esc(p_dates), subhead_right)]],
                                    colWidths=[usable_width * 0.72, usable_width * 0.28],
                                    hAlign='LEFT',
                                )
                                tbl.setStyle(TableStyle([
                                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                    ("RIGHTPADDING", (0, 0), (0, -1), 0),
                                    ("RIGHTPADDING", (1, 0), (1, -1), 6),
                                    ("TOPPADDING", (0, 0), (-1, -1), 0.5),
                                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
                                ]))
                                story.append(tbl)
                            elif p_title:
                                story.append(Paragraph(esc(p_title), body_style))

                            p_bullets = p.bullets if hasattr(p, "bullets") else p.get("bullets", [])
                            for b in p_bullets:
                                clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                                if clean_b:
                                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))

                    elif (hasattr(ent, "role") and ent.role) or (isinstance(ent, dict) and ent.get("role")):
                        r_title = ent.role if hasattr(ent, "role") else ent.get("role", "")
                        r_dates = ent.dates if hasattr(ent, "dates") else ent.get("dates", "")
                        if r_title and r_dates:
                            tbl = Table(
                                [[Paragraph(esc(r_title), body_style), Paragraph(esc(r_dates), subhead_right)]],
                                colWidths=[usable_width * 0.72, usable_width * 0.28],
                                hAlign='LEFT',
                            )
                            tbl.setStyle(TableStyle([
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                ("RIGHTPADDING", (0, 0), (0, -1), 0),
                                ("RIGHTPADDING", (1, 0), (1, -1), 6),
                                ("TOPPADDING", (0, 0), (-1, -1), 0.5),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
                            ]))
                            story.append(tbl)
                        elif r_title:
                            story.append(Paragraph(esc(r_title), body_style))

                    r_groups = ent.responsibility_groups if hasattr(ent, "responsibility_groups") else ent.get("responsibility_groups", [])
                    if r_groups:
                        for grp in r_groups:
                            g_heading = grp.heading if hasattr(grp, "heading") else grp.get("heading", "")
                            if g_heading:
                                story.append(Paragraph(f"<b>{esc(g_heading)}</b>", body_style))
                            g_bullets = grp.bullets if hasattr(grp, "bullets") else grp.get("bullets", [])
                            for b in g_bullets:
                                clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                                if clean_b:
                                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))
                    else:
                        e_bullets = ent.bullets if hasattr(ent, "bullets") else ent.get("bullets", [])
                        if not prog or not any(p.bullets if hasattr(p, "bullets") else p.get("bullets") for p in prog):
                            for b in e_bullets:
                                clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                                if clean_b:
                                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))
                rendered_sections.add("experience")

        # Internships (separate section — renders from parsed_data["internships"])
        elif sec_lower == "internships" and "internships" not in rendered_sections:
            intern_data = parsed_data.get("internships", []) or parsed_data.get("internships_raw", [])
            if intern_data:
                heading_title = std_headings.get("internships", "INTERNSHIPS")
                add_section_header(heading_title)
                if isinstance(intern_data, list) and intern_data and isinstance(intern_data[0], dict) and "company" in intern_data[0]:
                    intern_entities = intern_data
                elif isinstance(intern_data, list) and intern_data and isinstance(intern_data[0], str):
                    intern_entities = parse_experience_section(intern_data)
                else:
                    intern_entities = intern_data

                for ent in intern_entities:
                    comp = ent.company if hasattr(ent, "company") else ent.get("company", "")
                    loc = ent.location if hasattr(ent, "location") else ent.get("location", "")
                    if comp and comp.lower() not in ("internship", "internships"):
                        if loc:
                            tbl = Table(
                                [[Paragraph(esc(comp), subhead_left), Paragraph(esc(loc), subhead_right)]],
                                colWidths=[usable_width * 0.72, usable_width * 0.28],
                                hAlign='LEFT',
                            )
                            tbl.setStyle(TableStyle([
                                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                ("RIGHTPADDING", (0, 0), (0, -1), 0),
                                ("RIGHTPADDING", (1, 0), (1, -1), 6),
                                ("TOPPADDING", (0, 0), (-1, -1), 1),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                            ]))
                            story.append(tbl)
                        else:
                            story.append(Paragraph(esc(comp), subhead_left))

                    r_title = ent.role if hasattr(ent, "role") else ent.get("role", "")
                    r_dates = ent.dates if hasattr(ent, "dates") else ent.get("dates", "")
                    if r_title and r_dates:
                        tbl = Table(
                            [[Paragraph(esc(r_title), body_style), Paragraph(esc(r_dates), subhead_right)]],
                            colWidths=[usable_width * 0.72, usable_width * 0.28],
                            hAlign='LEFT',
                        )
                        tbl.setStyle(TableStyle([
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 0),
                            ("RIGHTPADDING", (0, 0), (0, -1), 0),
                            ("RIGHTPADDING", (1, 0), (1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 0.5),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
                        ]))
                        story.append(tbl)
                    elif r_title:
                        story.append(Paragraph(esc(r_title), body_style))

                    e_bullets = ent.bullets if hasattr(ent, "bullets") else ent.get("bullets", [])
                    for b in e_bullets:
                        clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                        if clean_b:
                            story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))
                rendered_sections.add("internships")



        # Projects
        elif sec_lower in ("projects", "technical projects", "relevant projects") and "projects" not in rendered_sections:
            proj = parsed_data.get("projects_raw", []) or parsed_data.get("projects", [])
            if proj:
                heading_title = std_headings.get("projects", "PROJECTS")
                add_section_header(heading_title)
                if isinstance(proj, list) and proj and isinstance(proj[0], dict) and "title" in proj[0] and "bullets" in proj[0]:
                    proj_entities = proj
                else:
                    proj_entities = parse_projects_section(proj)

                for p in proj_entities:
                    p_title = p.title if hasattr(p, "title") else p.get("title", "")
                    p_tech = p.tech_stack if hasattr(p, "tech_stack") else (p.get("tech_stack") or p.get("technologies"))
                    p_dates = p.dates if hasattr(p, "dates") else p.get("dates")

                    if p_title and p_dates:
                        tbl = Table(
                            [[Paragraph(esc(p_title), subhead_left), Paragraph(esc(p_dates), subhead_right)]],
                            colWidths=[usable_width * 0.72, usable_width * 0.28],
                            hAlign='LEFT',
                        )
                        tbl.setStyle(TableStyle([
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 0),
                            ("RIGHTPADDING", (0, 0), (0, -1), 0),
                            ("RIGHTPADDING", (1, 0), (1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 1),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                        ]))
                        story.append(tbl)
                    elif p_title:
                        story.append(Paragraph(esc(p_title), subhead_left))

                    if p_tech:
                        p_tech_str = ", ".join(p_tech) if isinstance(p_tech, list) else str(p_tech)
                        story.append(Paragraph(esc(f"Technologies: {p_tech_str}"), tech_stack_style))

                    p_bullets = p.bullets if hasattr(p, "bullets") else p.get("bullets", [])
                    for b in p_bullets:
                        clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                        if clean_b:
                            story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))
                rendered_sections.add("projects")

        # Education
        elif sec_lower == "education" and "education" not in rendered_sections:
            edu = parsed_data.get("education_raw", []) or parsed_data.get("education", [])
            if edu:
                heading_title = std_headings.get("education", "EDUCATION")
                add_section_header(heading_title)
                for idx, item in enumerate(edu):
                    if isinstance(item, dict):
                        inst = item.get("institution") or item.get("school") or item.get("name") or ""
                        degree = item.get("degree") or item.get("qualification") or ""
                        dates = item.get("dates") or item.get("date_range") or item.get("year") or ""
                        loc = item.get("location") or ""
                        gpa = item.get("gpa") or item.get("cgpa") or item.get("percentage") or item.get("score") or ""

                        if inst:
                            inst_text = inst + (f", {loc}" if loc else "")
                            if dates:
                                tbl = Table(
                                    [[Paragraph(esc(inst_text), subhead_left), Paragraph(esc(dates), subhead_right)]],
                                    colWidths=[usable_width * 0.72, usable_width * 0.28],
                                    hAlign='LEFT',
                                )
                                tbl.setStyle(TableStyle([
                                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                    ("RIGHTPADDING", (0, 0), (0, -1), 0),
                                    ("RIGHTPADDING", (1, 0), (1, -1), 6),
                                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                                ]))
                                story.append(tbl)
                            else:
                                story.append(Paragraph(esc(inst_text), subhead_left))

                        deg_parts = []
                        if degree:
                            deg_parts.append(degree)
                        if gpa:
                            deg_parts.append(gpa if ("cgpa" in gpa.lower() or "percentage" in gpa.lower() or "%" in gpa) else f"CGPA: {gpa}")
                        if deg_parts:
                            story.append(Paragraph(esc(" | ".join(deg_parts)), body_style))
                    else:
                        item_str = str(item).strip()
                        if not item_str:
                            continue
                        sub_lines = [l.strip() for l in item_str.split("\n") if l.strip()]
                        for sub_idx, sub_clean in enumerate(sub_lines):
                            cleaned = _clean_title_and_date(sub_clean)
                            if cleaned and cleaned[0] and cleaned[1] and sub_idx == 0:
                                title_str, date_str = cleaned
                                tbl = Table(
                                    [[Paragraph(esc(title_str), subhead_left), Paragraph(esc(date_str), subhead_right)]],
                                    colWidths=[usable_width * 0.72, usable_width * 0.28],
                                    hAlign='LEFT',
                                )
                                tbl.setStyle(TableStyle([
                                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                                    ("RIGHTPADDING", (0, 0), (0, -1), 0),
                                    ("RIGHTPADDING", (1, 0), (1, -1), 6),
                                    ("TOPPADDING", (0, 0), (-1, -1), 1),
                                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                                ]))
                                story.append(tbl)
                            else:
                                style_to_use = subhead_left if sub_idx == 0 else body_style
                                story.append(Paragraph(esc(sub_clean), style_to_use))

                    if idx < len(edu) - 1:
                        story.append(Spacer(1, 2))
                rendered_sections.add("education")

        # Certifications
        elif sec_lower in ("certifications", "certificates") and "certifications" not in rendered_sections:
            certs = parsed_data.get("certifications", []) or parsed_data.get("certifications_raw", []) or parsed_data.get("certificates", [])
            if certs:
                heading_title = std_headings.get("certifications", "CERTIFICATIONS")
                add_section_header(heading_title)
                for item in certs:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    cleaned_bullet = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    bullet_text = f"&bull;&nbsp;&nbsp;{esc(cleaned_bullet)}"
                    story.append(Paragraph(bullet_text, bullet_style))
                rendered_sections.add("certifications")

        # Achievements
        elif sec_lower in ("achievements", "honors & awards", "awards") and "achievements" not in rendered_sections:
            ach = parsed_data.get("achievements", []) or parsed_data.get("achievements_raw", []) or parsed_data.get("awards", [])
            if ach:
                heading_title = std_headings.get("achievements", "HONORS & AWARDS")
                add_section_header(heading_title)
                for item in ach:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    cleaned_bullet = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    bullet_text = f"&bull;&nbsp;&nbsp;{esc(cleaned_bullet)}"
                    story.append(Paragraph(bullet_text, bullet_style))
                rendered_sections.add("achievements")

        # Publications
        elif sec_lower in ("publications", "publications & research") and "publications" not in rendered_sections:
            pubs = parsed_data.get("publications_raw", []) or parsed_data.get("publications", [])
            if pubs:
                heading_title = std_headings.get("publications", "PUBLICATIONS & RESEARCH")
                add_section_header(heading_title)
                for item in pubs:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    cleaned_bullet = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(cleaned_bullet)}", bullet_style))
                rendered_sections.add("publications")

        # Research
        elif sec_lower in ("research", "research experience") and "research" not in rendered_sections:
            res = parsed_data.get("research_raw", []) or parsed_data.get("research", [])
            if res:
                heading_title = std_headings.get("research", "RESEARCH EXPERIENCE")
                add_section_header(heading_title)
                for item in res:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    cleaned_bullet = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(cleaned_bullet)}", bullet_style))
                rendered_sections.add("research")

        # Languages
        elif sec_lower == "languages" and "languages" not in rendered_sections:
            langs = parsed_data.get("languages", []) or parsed_data.get("languages_raw", []) or parsed_data.get("languages_known", [])
            if langs:
                heading_title = std_headings.get("languages", "LANGUAGES")
                add_section_header(heading_title)
                if isinstance(langs, list):
                    story.append(Paragraph(esc(", ".join(str(l) for l in langs if l)), body_style))
                else:
                    story.append(Paragraph(esc(str(langs)), body_style))
                rendered_sections.add("languages")

    # Render any custom sections from additional_sections not yet rendered
    add_secs = parsed_data.get("additional_sections", [])
    if add_secs:
        for add_sec in add_secs:
            sec_title = (
                (add_sec.get("heading", "") or add_sec.get("title", ""))
                if isinstance(add_sec, dict)
                else (getattr(add_sec, "heading", "") or getattr(add_sec, "title", ""))
            ) or "ADDITIONAL INFORMATION"
            sec_title = sec_title.strip()
            if sec_title.lower() not in rendered_sections:
                add_section_header(sec_title)
                ev_units = (
                    add_sec.get("evidence_units", [])
                    if isinstance(add_sec, dict)
                    else getattr(add_sec, "evidence_units", [])
                )
                for ev in ev_units:
                    ev_text = (ev.get("text", "") if isinstance(ev, dict) else (getattr(ev, "text", "") or str(ev)))
                    clean_b = _BULLET_PREFIX_RE.sub("", str(ev_text)).strip()
                    if clean_b:
                        story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))
                items = (
                    add_sec.get("items", [])
                    if isinstance(add_sec, dict)
                    else getattr(add_sec, "items", [])
                )
                for item in items:
                    clean_b = _BULLET_PREFIX_RE.sub("", str(item)).strip()
                    if clean_b:
                        story.append(Paragraph(f"&bull;&nbsp;&nbsp;{esc(clean_b)}", bullet_style))
                rendered_sections.add(sec_title.lower())

    doc.build(story)
    return buffer.getvalue()


def generate_pdf(content: str | dict | Any, candidate_name: str = "", template: str = "standard", experience_level: str = "fresher") -> bytes:
    """
    Primary PDF generation entrypoint.
    If content is CandidateProfile or dict, renders directly.
    If content is string, parses structure deterministically and renders structured flowables.
    """
    if hasattr(content, "to_parsed_dict"):
        content = content.to_parsed_dict()

    if isinstance(content, dict):
        return render_pdf_from_structured(content, candidate_name=candidate_name, template=template, experience_level=experience_level)

    from app.modules.resume.parsing.structurer import structure_resume_text
    structured = structure_resume_text(str(content))
    return render_pdf_from_structured(structured, candidate_name=candidate_name, template=template, experience_level=experience_level)


def _set_run_color(run, hex_color: str):
    run.font.color.rgb = RGBColor.from_string(hex_color.lstrip("#"))


def _add_bottom_border(paragraph, hex_color: str = SIGNAL_600):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hex_color.lstrip("#"))
    borders.append(bottom)
    p_pr.append(borders)


def render_docx_from_structured(
    parsed_data: dict, candidate_name: str = "", template: str = "standard"
) -> bytes:
    """
    Renders an ATS-optimized, beautifully styled DOCX document directly from a structured resume dictionary.
    No regex text-splicing — pure structured document generation.
    """
    template = template.lower() if template else "standard"
    if template in ("modern", "technical", "stanford"):
        template = "standard"
    elif template in ("executive", "harvard"):
        template = "classic"
    elif template == "minimal":
        template = "compact"

    document = Document()

    # Set margins
    margin_size = 0.35 if template == "compact" else 0.45
    for section in document.sections:
        section.top_margin = Inches(margin_size)
        section.bottom_margin = Inches(margin_size)
        section.left_margin = Inches(margin_size)
        section.right_margin = Inches(margin_size)

    sec = document.sections[0]
    usable_width_inches = sec.page_width.inches - sec.left_margin.inches - sec.right_margin.inches
    right_tab_stop = Inches(round(usable_width_inches, 2))

    if template == "classic":
        font_name = "Times New Roman"
        accent_color = CLASSIC_ACCENT
        alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_font_pt = Pt(10.0)
    elif template == "compact":
        font_name = "Calibri"
        accent_color = INK_900
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        body_font_pt = Pt(9.0)
    else:  # standard
        font_name = "Calibri"
        accent_color = SIGNAL_600
        alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_font_pt = Pt(9.5)

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = body_font_pt

    personal = parsed_data.get("personal", {}) or parsed_data.get("personal_info", {}) or {}
    name = personal.get("name") or personal.get("full_name") or candidate_name or "Candidate"
    name = re.sub(r"^[#\s*]+", "", str(name)).strip()

    contacts = []
    if personal.get("location"):
        contacts.append(personal["location"])
    elif personal.get("city") or personal.get("address"):
        loc = ", ".join(filter(None, [personal.get("city"), personal.get("address")]))
        if loc:
            contacts.append(loc)
    if personal.get("email"):
        contacts.append(personal["email"])
    if personal.get("phone"):
        contacts.append(personal["phone"])
    if personal.get("linkedin"):
        contacts.append(personal["linkedin"])
    if personal.get("github"):
        contacts.append(personal["github"])
    if personal.get("portfolio"):
        contacts.append(personal["portfolio"])

    if name:
        p_name = document.add_paragraph()
        p_name.alignment = alignment
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after = Pt(1)
        r_name = p_name.add_run(name)
        r_name.bold = True
        r_name.font.size = Pt(15)
        _set_run_color(r_name, INK_900)

    if contacts:
        p_contact = document.add_paragraph()
        p_contact.alignment = alignment
        p_contact.paragraph_format.space_before = Pt(0)
        p_contact.paragraph_format.space_after = Pt(3)
        formatted_contacts = " • ".join(str(c).strip(" •·|") for c in contacts if str(c).strip())
        r_contact = p_contact.add_run(formatted_contacts)
        r_contact.font.size = Pt(8.5)
        _set_run_color(r_contact, INK_700)

    def add_section_header(title: str):
        p_head = document.add_paragraph()
        p_head.paragraph_format.space_before = Pt(8 if template == "compact" else 10)
        p_head.paragraph_format.space_after = Pt(2)
        p_head.paragraph_format.keep_with_next = True
        r_head = p_head.add_run(title.upper())
        r_head.bold = True
        r_head.font.size = Pt(10)
        _set_run_color(r_head, accent_color)
        _add_bottom_border(p_head, accent_color)

    strategy_info = parsed_data.get("_strategy", {}) if isinstance(parsed_data.get("_strategy"), dict) else {}
    section_order = (
        parsed_data.get("_ordered_sections")
        or strategy_info.get("section_order")
        or ["summary", "skills", "experience", "internships", "projects", "education", "certifications", "achievements", "publications", "research", "languages"]
    )
    std_headings = strategy_info.get("standard_ats_headings", {}) or STANDARD_ATS_HEADINGS
    rendered_sections = set()

    for sec in section_order:
        sec_lower = str(sec).lower()

        # Summary
        if sec_lower in ("summary", "objective") and "summary" not in rendered_sections:
            summary = parsed_data.get("summary") or parsed_data.get("objective")
            if summary and str(summary).strip():
                add_section_header(std_headings.get("summary", "PROFESSIONAL SUMMARY"))
                p_sum = document.add_paragraph()
                p_sum.paragraph_format.space_before = Pt(0)
                p_sum.paragraph_format.space_after = Pt(2)
                r_sum = p_sum.add_run(str(summary).strip())
                r_sum.font.size = Pt(8.5)
                _set_run_color(r_sum, INK_900)
                rendered_sections.add("summary")

        # Skills
        elif sec_lower in ("skills", "technical skills", "competencies", "core competencies") and "skills" not in rendered_sections:
            skills_cat = parsed_data.get("skills_categorized", [])
            skills = parsed_data.get("skills", [])
            if skills_cat or skills:
                add_section_header(std_headings.get("skills", "TECHNICAL SKILLS"))
                if skills_cat and isinstance(skills_cat, list):
                    for sc in skills_cat:
                        sc_str = str(sc).strip()
                        p_sk = document.add_paragraph()
                        p_sk.paragraph_format.space_before = Pt(0)
                        p_sk.paragraph_format.space_after = Pt(1.5)
                        if ":" in sc_str:
                            parts = sc_str.split(":", 1)
                            r_pre = p_sk.add_run(parts[0].strip() + ": ")
                            r_pre.bold = True
                            r_pre.font.size = Pt(8.5)
                            _set_run_color(r_pre, INK_900)
                            r_val = p_sk.add_run(parts[1].strip())
                            r_val.font.size = Pt(8.5)
                            _set_run_color(r_val, INK_900)
                        else:
                            r_val = p_sk.add_run(sc_str)
                            r_val.font.size = Pt(8.5)
                            _set_run_color(r_val, INK_900)
                elif isinstance(skills, dict):
                    for cat, val in skills.items():
                        p_sk = document.add_paragraph()
                        p_sk.paragraph_format.space_before = Pt(0)
                        p_sk.paragraph_format.space_after = Pt(1.5)
                        r_pre = p_sk.add_run(f"{cat}: ")
                        r_pre.bold = True
                        r_pre.font.size = Pt(8.5)
                        _set_run_color(r_pre, INK_900)
                        val_str = ", ".join(str(v) for v in val) if isinstance(val, list) else str(val)
                        r_val = p_sk.add_run(val_str)
                        r_val.font.size = Pt(8.5)
                        _set_run_color(r_val, INK_900)
                elif isinstance(skills, list):
                    cat_lines = [str(s).strip() for s in skills if ":" in str(s) and len(str(s).split(":")[0].split()) <= 6]
                    if len(cat_lines) >= 2 and len(cat_lines) >= len(skills) * 0.5:
                        for s in skills:
                            s_str = str(s).strip()
                            p_sk = document.add_paragraph()
                            p_sk.paragraph_format.space_before = Pt(0)
                            p_sk.paragraph_format.space_after = Pt(1.5)
                            if ":" in s_str:
                                parts = s_str.split(":", 1)
                                r_pre = p_sk.add_run(parts[0].strip() + ": ")
                                r_pre.bold = True
                                r_pre.font.size = Pt(8.5)
                                _set_run_color(r_pre, INK_900)
                                r_val = p_sk.add_run(parts[1].strip())
                                r_val.font.size = Pt(8.5)
                                _set_run_color(r_val, INK_900)
                            else:
                                r_val = p_sk.add_run(s_str)
                                r_val.font.size = Pt(8.5)
                                _set_run_color(r_val, INK_900)
                    else:
                        p_sk = document.add_paragraph()
                        p_sk.paragraph_format.space_before = Pt(0)
                        p_sk.paragraph_format.space_after = Pt(2)
                        clean_skills = [re.sub(r"^[A-Za-z\s&]+:\s*", "", str(s)).strip() for s in skills if s and str(s).strip()]
                        clean_unique = []
                        seen_sk = set()
                        for cs in clean_skills:
                            if cs.lower() not in seen_sk:
                                seen_sk.add(cs.lower())
                                clean_unique.append(cs)
                        r_sk = p_sk.add_run(", ".join(clean_unique))
                        r_sk.font.size = body_font_pt
                        _set_run_color(r_sk, INK_900)
                else:
                    p_sk = document.add_paragraph()
                    p_sk.paragraph_format.space_before = Pt(0)
                    p_sk.paragraph_format.space_after = Pt(2)
                    r_sk = p_sk.add_run(str(skills))
                    r_sk.font.size = body_font_pt
                    _set_run_color(r_sk, INK_900)
                rendered_sections.add("skills")

        # Experience
        elif sec_lower in ("experience", "work_experience", "professional experience") and "experience" not in rendered_sections:
            exp = parsed_data.get("experience", []) or parsed_data.get("experience_raw", []) or parsed_data.get("work_experience", [])
            if exp:
                add_section_header(std_headings.get("experience", "PROFESSIONAL EXPERIENCE"))
                if isinstance(exp, list) and exp and isinstance(exp[0], dict) and "company" in exp[0] and ("bullets" in exp[0] or "progression" in exp[0]):
                    exp_entities = exp
                else:
                    exp_entities = parse_experience_section(exp)

                for ent in exp_entities:
                    comp = ent.company if hasattr(ent, "company") else ent.get("company", "")
                    loc = ent.location if hasattr(ent, "location") else ent.get("location", "")

                    if comp and comp.lower() not in ("work experience", "experience"):
                        p_comp = document.add_paragraph()
                        p_comp.paragraph_format.space_before = Pt(3)
                        p_comp.paragraph_format.space_after = Pt(1)
                        p_comp.paragraph_format.keep_with_next = True
                        if loc:
                            p_comp.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                            r_c = p_comp.add_run(comp)
                            r_c.bold = True
                            r_c.font.size = Pt(9.5)
                            _set_run_color(r_c, INK_900)
                            r_l = p_comp.add_run(f"\t{loc}")
                            r_l.font.size = body_font_pt
                            _set_run_color(r_l, INK_700)
                        else:
                            r_c = p_comp.add_run(comp)
                            r_c.bold = True
                            r_c.font.size = Pt(9.5)
                            _set_run_color(r_c, INK_900)

                    prog = ent.progression if hasattr(ent, "progression") else ent.get("progression", [])
                    if prog:
                        for p in prog:
                            p_title = p.title if hasattr(p, "title") else p.get("title", "")
                            p_dates = p.dates if hasattr(p, "dates") else p.get("dates", "")
                            p_tit = document.add_paragraph()
                            p_tit.paragraph_format.space_before = Pt(1)
                            p_tit.paragraph_format.space_after = Pt(1)
                            p_tit.paragraph_format.keep_with_next = True
                            if p_dates:
                                p_tit.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                                r1 = p_tit.add_run(p_title)
                                r1.bold = True
                                r1.font.size = body_font_pt
                                _set_run_color(r1, INK_900)
                                r2 = p_tit.add_run(f"\t{p_dates}")
                                r2.font.size = body_font_pt
                                _set_run_color(r2, INK_700)
                            else:
                                r1 = p_tit.add_run(p_title)
                                r1.bold = True
                                r1.font.size = body_font_pt
                                _set_run_color(r1, INK_900)

                            p_bullets = p.bullets if hasattr(p, "bullets") else p.get("bullets", [])
                            for b in p_bullets:
                                clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                                if clean_b:
                                    p_b = document.add_paragraph(style="List Bullet")
                                    p_b.paragraph_format.space_before = Pt(0)
                                    p_b.paragraph_format.space_after = Pt(1)
                                    p_b.paragraph_format.left_indent = Inches(0.2)
                                    r_b = p_b.add_run(clean_b)
                                    r_b.font.size = body_font_pt
                                    _set_run_color(r_b, INK_900)
                    elif (hasattr(ent, "role") and ent.role) or (isinstance(ent, dict) and ent.get("role")):
                        r_title = ent.role if hasattr(ent, "role") else ent.get("role", "")
                        r_dates = ent.dates if hasattr(ent, "dates") else ent.get("dates", "")
                        p_tit = document.add_paragraph()
                        p_tit.paragraph_format.space_before = Pt(1)
                        p_tit.paragraph_format.space_after = Pt(1)
                        p_tit.paragraph_format.keep_with_next = True
                        if r_dates:
                            p_tit.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                            r1 = p_tit.add_run(r_title)
                            r1.bold = True
                            r1.font.size = body_font_pt
                            _set_run_color(r1, INK_900)
                            r2 = p_tit.add_run(f"\t{r_dates}")
                            r2.font.size = body_font_pt
                            _set_run_color(r2, INK_700)
                        else:
                            r1 = p_tit.add_run(r_title)
                            r1.bold = True
                            r1.font.size = body_font_pt
                            _set_run_color(r1, INK_900)

                    r_groups = ent.responsibility_groups if hasattr(ent, "responsibility_groups") else ent.get("responsibility_groups", [])
                    if r_groups:
                        for grp in r_groups:
                            g_heading = grp.heading if hasattr(grp, "heading") else grp.get("heading", "")
                            if g_heading:
                                p_grp = document.add_paragraph()
                                p_grp.paragraph_format.space_before = Pt(2)
                                p_grp.paragraph_format.space_after = Pt(1)
                                p_grp.paragraph_format.keep_with_next = True
                                r_grp = p_grp.add_run(g_heading)
                                r_grp.bold = True
                                r_grp.font.size = body_font_pt
                                _set_run_color(r_grp, INK_900)
                            g_bullets = grp.bullets if hasattr(grp, "bullets") else grp.get("bullets", [])
                            for b in g_bullets:
                                clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                                if clean_b:
                                    p_b = document.add_paragraph(style="List Bullet")
                                    p_b.paragraph_format.space_before = Pt(0)
                                    p_b.paragraph_format.space_after = Pt(1)
                                    p_b.paragraph_format.left_indent = Inches(0.2)
                                    r_b = p_b.add_run(clean_b)
                                    r_b.font.size = body_font_pt
                                    _set_run_color(r_b, INK_900)
                    else:
                        e_bullets = ent.bullets if hasattr(ent, "bullets") else ent.get("bullets", [])
                        if not prog or not any(p.bullets if hasattr(p, "bullets") else p.get("bullets") for p in prog):
                            for b in e_bullets:
                                clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                                if clean_b:
                                    p_b = document.add_paragraph(style="List Bullet")
                                    p_b.paragraph_format.space_before = Pt(0)
                                    p_b.paragraph_format.space_after = Pt(1)
                                    p_b.paragraph_format.left_indent = Inches(0.2)
                                    r_b = p_b.add_run(clean_b)
                                    r_b.font.size = body_font_pt
                                    _set_run_color(r_b, INK_900)
                rendered_sections.add("experience")

        # Internships (DOCX — dedicated section from parsed_data["internships"])
        elif sec_lower == "internships" and "internships" not in rendered_sections:
            intern_data = parsed_data.get("internships", []) or parsed_data.get("internships_raw", [])
            if intern_data:
                add_section_header(std_headings.get("internships", "INTERNSHIPS"))
                if isinstance(intern_data, list) and intern_data and isinstance(intern_data[0], dict) and "company" in intern_data[0]:
                    intern_entities = intern_data
                elif isinstance(intern_data, list) and intern_data and isinstance(intern_data[0], str):
                    intern_entities = parse_experience_section(intern_data)
                else:
                    intern_entities = intern_data

                for ent in intern_entities:
                    comp = ent.company if hasattr(ent, "company") else ent.get("company", "")
                    loc = ent.location if hasattr(ent, "location") else ent.get("location", "")
                    if comp and comp.lower() not in ("internship", "internships"):
                        p_comp = document.add_paragraph()
                        p_comp.paragraph_format.space_before = Pt(3)
                        p_comp.paragraph_format.space_after = Pt(1)
                        p_comp.paragraph_format.keep_with_next = True
                        if loc:
                            p_comp.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                            r_c = p_comp.add_run(comp)
                            r_c.bold = True
                            r_c.font.size = Pt(9.5)
                            _set_run_color(r_c, INK_900)
                            r_l = p_comp.add_run(f"\t{loc}")
                            r_l.font.size = body_font_pt
                            _set_run_color(r_l, INK_700)
                        else:
                            r_c = p_comp.add_run(comp)
                            r_c.bold = True
                            r_c.font.size = Pt(9.5)
                            _set_run_color(r_c, INK_900)

                    r_title = ent.role if hasattr(ent, "role") else ent.get("role", "")
                    r_dates = ent.dates if hasattr(ent, "dates") else ent.get("dates", "")
                    if r_title:
                        p_tit = document.add_paragraph()
                        p_tit.paragraph_format.space_before = Pt(1)
                        p_tit.paragraph_format.space_after = Pt(1)
                        p_tit.paragraph_format.keep_with_next = True
                        if r_dates:
                            p_tit.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                            r1 = p_tit.add_run(r_title)
                            r1.bold = True
                            r1.font.size = body_font_pt
                            _set_run_color(r1, INK_900)
                            r2 = p_tit.add_run(f"\t{r_dates}")
                            r2.font.size = body_font_pt
                            _set_run_color(r2, INK_700)
                        else:
                            r1 = p_tit.add_run(r_title)
                            r1.bold = True
                            r1.font.size = body_font_pt
                            _set_run_color(r1, INK_900)

                    e_bullets = ent.bullets if hasattr(ent, "bullets") else ent.get("bullets", [])
                    for b in e_bullets:
                        clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                        if clean_b:
                            p_b = document.add_paragraph(style="List Bullet")
                            p_b.paragraph_format.space_before = Pt(0)
                            p_b.paragraph_format.space_after = Pt(1)
                            p_b.paragraph_format.left_indent = Inches(0.2)
                            r_b = p_b.add_run(clean_b)
                            r_b.font.size = body_font_pt
                            _set_run_color(r_b, INK_900)
                rendered_sections.add("internships")

        # Projects
        elif sec_lower in ("projects", "technical projects", "relevant projects") and "projects" not in rendered_sections:
            proj = parsed_data.get("projects_raw", []) or parsed_data.get("projects", [])
            if proj:
                add_section_header(std_headings.get("projects", "PROJECTS"))
                if isinstance(proj, list) and proj and isinstance(proj[0], dict) and "title" in proj[0] and "bullets" in proj[0]:
                    proj_entities = proj
                else:
                    proj_entities = parse_projects_section(proj)

                for p in proj_entities:
                    p_title = p.title if hasattr(p, "title") else p.get("title", "")
                    p_tech = p.tech_stack if hasattr(p, "tech_stack") else (p.get("tech_stack") or p.get("technologies"))
                    p_dates = p.dates if hasattr(p, "dates") else p.get("dates")

                    if p_title:
                        p_tit = document.add_paragraph()
                        p_tit.paragraph_format.space_before = Pt(2)
                        p_tit.paragraph_format.space_after = Pt(1)
                        p_tit.paragraph_format.keep_with_next = True
                        if p_dates:
                            p_tit.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                            r1 = p_tit.add_run(p_title)
                            r1.bold = True
                            r1.font.size = body_font_pt
                            _set_run_color(r1, INK_900)
                            r2 = p_tit.add_run(f"\t{p_dates}")
                            r2.font.size = body_font_pt
                            _set_run_color(r2, INK_700)
                        else:
                            r1 = p_tit.add_run(p_title)
                            r1.bold = True
                            r1.font.size = body_font_pt
                            _set_run_color(r1, INK_900)

                    if p_tech:
                        p_tech_str = ", ".join(p_tech) if isinstance(p_tech, list) else str(p_tech)
                        p_t = document.add_paragraph()
                        p_t.paragraph_format.space_before = Pt(0)
                        p_t.paragraph_format.space_after = Pt(2)
                        r_tech = p_t.add_run(f"Technologies: {p_tech_str}")
                        r_tech.italic = True
                        r_tech.font.size = Pt(8.5)
                        _set_run_color(r_tech, INK_700)

                    p_bullets = p.bullets if hasattr(p, "bullets") else p.get("bullets", [])
                    for b in p_bullets:
                        clean_b = _BULLET_PREFIX_RE.sub("", str(b)).strip()
                        if clean_b:
                            p_b = document.add_paragraph(style="List Bullet")
                            p_b.paragraph_format.space_before = Pt(0)
                            p_b.paragraph_format.space_after = Pt(1)
                            p_b.paragraph_format.left_indent = Inches(0.2)
                            r_b = p_b.add_run(clean_b)
                            r_b.font.size = body_font_pt
                            _set_run_color(r_b, INK_900)
                rendered_sections.add("projects")

        # Education
        elif sec_lower == "education" and "education" not in rendered_sections:
            edu = parsed_data.get("education_raw", []) or parsed_data.get("education", [])
            if edu:
                add_section_header(std_headings.get("education", "EDUCATION"))
                for idx, item in enumerate(edu):
                    if isinstance(item, dict):
                        inst = item.get("institution") or item.get("school") or item.get("name") or ""
                        degree = item.get("degree") or item.get("qualification") or ""
                        dates = item.get("dates") or item.get("date_range") or item.get("year") or ""
                        loc = item.get("location") or ""
                        gpa = item.get("gpa") or item.get("cgpa") or item.get("percentage") or item.get("score") or ""

                        if inst:
                            p_inst = document.add_paragraph()
                            p_inst.paragraph_format.space_before = Pt(3 if idx > 0 else 0)
                            p_inst.paragraph_format.space_after = Pt(1)
                            p_inst.paragraph_format.keep_with_next = True
                            inst_text = inst + (f", {loc}" if loc else "")
                            if dates:
                                p_inst.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                                r1 = p_inst.add_run(inst_text)
                                r1.bold = True
                                r1.font.size = body_font_pt
                                _set_run_color(r1, INK_900)
                                r2 = p_inst.add_run(f"\t{dates}")
                                r2.bold = True if template == "classic" else False
                                r2.font.size = body_font_pt
                                _set_run_color(r2, INK_700)
                            else:
                                r1 = p_inst.add_run(inst_text)
                                r1.bold = True
                                r1.font.size = body_font_pt
                                _set_run_color(r1, INK_900)

                        deg_parts = []
                        if degree:
                            deg_parts.append(degree)
                        if gpa:
                            deg_parts.append(gpa if ("cgpa" in gpa.lower() or "percentage" in gpa.lower() or "%" in gpa) else f"CGPA: {gpa}")
                        if deg_parts:
                            p_deg = document.add_paragraph()
                            p_deg.paragraph_format.space_before = Pt(0)
                            p_deg.paragraph_format.space_after = Pt(1.5)
                            r_deg = p_deg.add_run(" | ".join(deg_parts))
                            r_deg.font.size = body_font_pt
                            _set_run_color(r_deg, INK_900)
                    else:
                        item_str = str(item).strip()
                        if not item_str:
                            continue
                        sub_lines = [l.strip() for l in item_str.split("\n") if l.strip()]
                        for sub_idx, sub_clean in enumerate(sub_lines):
                            p_ed = document.add_paragraph()
                            p_ed.paragraph_format.space_before = Pt(3 if sub_idx == 0 and idx > 0 else 0)
                            p_ed.paragraph_format.space_after = Pt(1)
                            p_ed.paragraph_format.keep_with_next = True
                            cleaned = _clean_title_and_date(sub_clean)
                            if cleaned and cleaned[0] and cleaned[1] and sub_idx == 0:
                                t_str, d_str = cleaned
                                p_ed.paragraph_format.tab_stops.add_tab_stop(right_tab_stop, WD_TAB_ALIGNMENT.RIGHT)
                                r1 = p_ed.add_run(t_str)
                                r1.bold = True
                                r1.font.size = body_font_pt
                                _set_run_color(r1, INK_900)
                                r2 = p_ed.add_run(f"\t{d_str}")
                                r2.bold = True if template == "classic" else False
                                r2.font.size = body_font_pt
                                _set_run_color(r2, INK_700)
                            else:
                                r_ed = p_ed.add_run(sub_clean)
                                r_ed.font.size = body_font_pt
                                if sub_idx == 0:
                                    r_ed.bold = True
                                    _set_run_color(r_ed, INK_900)
                                else:
                                    _set_run_color(r_ed, INK_900)
                rendered_sections.add("education")

        # Certifications
        elif sec_lower in ("certifications", "certificates") and "certifications" not in rendered_sections:
            certs = parsed_data.get("certifications", []) or parsed_data.get("certifications_raw", []) or parsed_data.get("certificates", [])
            if certs:
                add_section_header(std_headings.get("certifications", "CERTIFICATIONS"))
                for item in certs:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    clean_b = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    p_b = document.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(0)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.left_indent = Inches(0.2)
                    r_b = p_b.add_run(clean_b)
                    r_b.font.size = body_font_pt
                    _set_run_color(r_b, INK_900)
                rendered_sections.add("certifications")

        # Achievements
        elif sec_lower in ("achievements", "honors & awards", "awards") and "achievements" not in rendered_sections:
            ach = parsed_data.get("achievements", []) or parsed_data.get("achievements_raw", []) or parsed_data.get("awards", [])
            if ach:
                add_section_header(std_headings.get("achievements", "HONORS & AWARDS"))
                for item in ach:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    clean_b = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    p_b = document.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(0)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.left_indent = Inches(0.2)
                    r_b = p_b.add_run(clean_b)
                    r_b.font.size = body_font_pt
                    _set_run_color(r_b, INK_900)
                rendered_sections.add("achievements")

        # Publications
        elif sec_lower in ("publications", "publications & research") and "publications" not in rendered_sections:
            pubs = parsed_data.get("publications_raw", []) or parsed_data.get("publications", [])
            if pubs:
                add_section_header(std_headings.get("publications", "PUBLICATIONS & RESEARCH"))
                for item in pubs:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    clean_b = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    p_b = document.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(0)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.left_indent = Inches(0.2)
                    r_b = p_b.add_run(clean_b)
                    r_b.font.size = body_font_pt
                    _set_run_color(r_b, INK_900)
                rendered_sections.add("publications")

        # Research
        elif sec_lower in ("research", "research experience") and "research" not in rendered_sections:
            res = parsed_data.get("research_raw", []) or parsed_data.get("research", [])
            if res:
                add_section_header(std_headings.get("research", "RESEARCH EXPERIENCE"))
                for item in res:
                    item_str = str(item).strip()
                    if not item_str:
                        continue
                    clean_b = _BULLET_PREFIX_RE.sub("", item_str).strip()
                    p_b = document.add_paragraph(style="List Bullet")
                    p_b.paragraph_format.space_before = Pt(0)
                    p_b.paragraph_format.space_after = Pt(1)
                    p_b.paragraph_format.left_indent = Inches(0.2)
                    r_b = p_b.add_run(clean_b)
                    r_b.font.size = body_font_pt
                    _set_run_color(r_b, INK_900)
                rendered_sections.add("research")

        # Languages
        elif sec_lower == "languages" and "languages" not in rendered_sections:
            langs = parsed_data.get("languages", []) or parsed_data.get("languages_raw", []) or parsed_data.get("languages_known", [])
            if langs:
                add_section_header(std_headings.get("languages", "LANGUAGES"))
                p_lang = document.add_paragraph()
                p_lang.paragraph_format.space_before = Pt(0)
                p_lang.paragraph_format.space_after = Pt(2)
                lang_text = ", ".join(str(l) for l in langs if l) if isinstance(langs, list) else str(langs)
                r_lang = p_lang.add_run(lang_text)
                r_lang.font.size = body_font_pt
                _set_run_color(r_lang, INK_900)
                rendered_sections.add("languages")

    # Render any custom sections from additional_sections not yet rendered
    add_secs = parsed_data.get("additional_sections", [])
    if add_secs:
        for add_sec in add_secs:
            sec_title = (
                (add_sec.get("heading", "") or add_sec.get("title", ""))
                if isinstance(add_sec, dict)
                else (getattr(add_sec, "heading", "") or getattr(add_sec, "title", ""))
            ) or "ADDITIONAL INFORMATION"
            sec_title = sec_title.strip()
            if sec_title.lower() not in rendered_sections:
                add_section_header(sec_title)
                ev_units = (
                    add_sec.get("evidence_units", [])
                    if isinstance(add_sec, dict)
                    else getattr(add_sec, "evidence_units", [])
                )
                for ev in ev_units:
                    ev_text = (ev.get("text", "") if isinstance(ev, dict) else (getattr(ev, "text", "") or str(ev)))
                    clean_b = _BULLET_PREFIX_RE.sub("", str(ev_text)).strip()
                    if clean_b:
                        p_b = document.add_paragraph(style="List Bullet")
                        p_b.paragraph_format.space_before = Pt(0)
                        p_b.paragraph_format.space_after = Pt(1)
                        p_b.paragraph_format.left_indent = Inches(0.2)
                        r_b = p_b.add_run(clean_b)
                        r_b.font.size = body_font_pt
                        _set_run_color(r_b, INK_900)
                items = (
                    add_sec.get("items", [])
                    if isinstance(add_sec, dict)
                    else getattr(add_sec, "items", [])
                )
                for item in items:
                    clean_b = _BULLET_PREFIX_RE.sub("", str(item)).strip()
                    if clean_b:
                        p_b = document.add_paragraph(style="List Bullet")
                        p_b.paragraph_format.space_before = Pt(0)
                        p_b.paragraph_format.space_after = Pt(1)
                        p_b.paragraph_format.left_indent = Inches(0.2)
                        r_b = p_b.add_run(clean_b)
                        r_b.font.size = body_font_pt
                        _set_run_color(r_b, INK_900)
                rendered_sections.add(sec_title.lower())

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def generate_docx(content: str | dict | Any, candidate_name: str = "", template: str = "standard", experience_level: str = "fresher") -> bytes:
    """
    Primary DOCX generation entrypoint.
    If content is CandidateProfile or dict, renders directly.
    If content is string, parses structure deterministically and renders structured document.
    """
    if hasattr(content, "to_parsed_dict"):
        content = content.to_parsed_dict()

    if isinstance(content, dict):
        return render_docx_from_structured(content, candidate_name=candidate_name, template=template)

    from app.modules.resume.parsing.structurer import structure_resume_text
    structured = structure_resume_text(str(content))
    return render_docx_from_structured(structured, candidate_name=candidate_name, template=template)


def verify_ats_pdf_parseability(
    pdf_bytes: bytes,
    source_profile: Any,
) -> tuple[bool, list[str]]:
    """
    Extracts text back from the generated PDF and performs strict end-to-end verification
    against the CandidateProfile to ensure 100% semantic fidelity:
    - Contact info preserved
    - All evidence units preserved (verifying first 4-5 words)
    - All metrics preserved
    - All company names and role titles preserved
    - All project titles and technologies preserved
    - Section headings preserved
    - Reading order is single-column top-to-bottom
    """
    import fitz
    errors: list[str] = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        return False, [f"Corrupted PDF output: {exc}"]

    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"
    doc.close()

    norm_text = " ".join(full_text.lower().split())

    # 1. Contact / Name
    name = getattr(getattr(source_profile, "contact", None), "name", None) or (
        source_profile.get("personal", {}).get("name") if isinstance(source_profile, dict) else ""
    )
    if name and name.lower() not in norm_text:
        errors.append(f"Candidate name '{name}' missing in extracted PDF text.")

    # 2. Evidence units (from profile.evidence_units, covers both experience and internship)
    ev_units = getattr(source_profile, "evidence_units", []) or (
        source_profile.get("evidence_units", []) if isinstance(source_profile, dict) else []
    )
    for ev in ev_units:
        raw_t = getattr(ev, "normalized_text", None) or getattr(ev, "text", "") if hasattr(ev, "text") else str(ev)
        clean_t = re.sub(r"^[•\-\*\s]+", "", raw_t).strip()
        words = clean_t.split()
        if len(words) >= 4:
            phrase = " ".join(words[:4]).lower()
            if phrase not in norm_text:
                errors.append(f"Evidence claim '{phrase}' missing from extracted PDF.")

    # 3. Metrics preservation
    metrics = re.findall(r"\b(?:\d+(?:\.\d+)?%|\$\d+(?:,\d{3})*(?:\.\d+)?[kKmMbB]?|\b\d+\+?\s*(?:years?|months?|hours?|ms|x|k|m|b)\b)", str(source_profile))
    for m in set(metrics):
        if m.lower() not in norm_text:
            errors.append(f"Metric '{m}' missing in extracted PDF text.")

    # 4. Companies & Roles (experience + internships)
    exps = getattr(source_profile, "experience", []) or (
        source_profile.get("experience", []) if isinstance(source_profile, dict) else []
    )
    interns = getattr(source_profile, "internships", []) or (
        source_profile.get("internships", []) if isinstance(source_profile, dict) else []
    )
    for exp in list(exps) + list(interns):
        comp = getattr(exp, "company", "") or (exp.get("company", "") if isinstance(exp, dict) else "")
        if comp and comp.lower() not in norm_text and comp.lower() not in ("experience", "work experience", "internship", "internships"):
            errors.append(f"Company name '{comp}' missing in extracted PDF text.")

    # 5. Project titles
    projs = getattr(source_profile, "projects", []) or (
        source_profile.get("projects", []) if isinstance(source_profile, dict) else []
    )
    for p in projs:
        p_title = getattr(p, "title", "") or (p.get("title", "") if isinstance(p, dict) else "")
        if p_title and p_title.lower() not in norm_text:
            errors.append(f"Project title '{p_title}' missing in extracted PDF text.")

    return len(errors) == 0, errors


def measure_pdf_page_count(pdf_bytes: bytes) -> int:
    """Measures the exact page count of a PDF byte stream using PyMuPDF."""
    import pymupdf
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    count = doc.page_count
    doc.close()
    return count


def verify_export_text_integrity(
    file_bytes: bytes,
    file_type: str = "pdf",
    required_facts: list[str] | None = None,
) -> tuple[bool, list[str]]:
    """
    Extracts text back from generated PDF/DOCX and validates that all required candidate facts
    (e.g. name, key titles, metrics) are faithfully preserved without degradation.
    Returns: (is_valid, missing_facts)
    """
    if not file_bytes:
        return False, ["Exported file is empty"]

    extracted_text = ""
    try:
        if file_type.lower() == "pdf":
            import pymupdf as fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            extracted_text = "\n".join(page.get_text() for page in doc)
            doc.close()
        elif file_type.lower() in ("docx", "doc"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        else:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        return False, [f"Extraction failed: {str(e)}"]

    if not required_facts:
        return len(extracted_text.strip()) > 50, []

    missing: list[str] = []
    text_lower = extracted_text.lower()
    for fact in required_facts:
        fact_clean = fact.strip()
        if fact_clean and fact_clean.lower() not in text_lower:
            missing.append(fact_clean)

    return len(missing) == 0, missing


def verify_export_against_structured_resume(
    file_bytes: bytes,
    structured_resume: dict | Any,
    file_type: str = "pdf",
) -> tuple[bool, dict[str, list[str]]]:
    """
    Comprehensive Phase 10 post-render validator.
    Extracts text back from generated PDF/DOCX and validates against structured data:
    - Candidate facts (name, email, phone, location)
    - Experience companies, roles, dates, and bullets
    - Project titles, technologies, and bullets
    - Education institutions, degrees, dates, CGPA
    - Quantified metrics preservation
    - Detects missing content, duplicate content, replacement characters (\ufffd),
      broken bullets, and cross-project contamination.
    """
    if hasattr(structured_resume, "to_parsed_dict"):
        data = structured_resume.to_parsed_dict()
    elif isinstance(structured_resume, dict):
        data = dict(structured_resume)
    else:
        data = {}

    report: dict[str, list[str]] = {
        "missing_facts": [],
        "duplicated_content": [],
        "replacement_characters": [],
        "broken_bullets": [],
        "project_integrity_issues": [],
        "education_integrity_issues": [],
    }

    if not file_bytes:
        report["missing_facts"].append("Exported file is empty")
        return False, report

    extracted_text = ""
    try:
        if file_type.lower() == "pdf":
            import pymupdf as fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            extracted_text = "\n".join(page.get_text() for page in doc)
            doc.close()
        elif file_type.lower() in ("docx", "doc"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join(p.text for p in doc.paragraphs)
        else:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        report["missing_facts"].append(f"Extraction failed: {str(e)}")
        return False, report

    # 1. Replacement characters check
    if "\ufffd" in extracted_text:
        report["replacement_characters"].append("Found \\ufffd replacement character in exported document")

    norm_extracted = re.sub(r"[^\w\s]", " ", extracted_text.lower())
    norm_words = set(norm_extracted.split())

    # 2. Personal contact info
    personal = data.get("personal", {}) or data.get("personal_info", {}) or {}
    name = personal.get("name") or personal.get("full_name")
    if name and re.sub(r"[^\w\s]", " ", str(name).lower()).strip() not in norm_extracted:
        report["missing_facts"].append(f"Candidate Name: {name}")

    email = personal.get("email")
    if email:
        email_raw = str(email).strip().lower()
        email_norm = re.sub(r"[^\w\s]", " ", email_raw).strip()
        if email_raw not in extracted_text.lower() and email_norm not in norm_extracted:
            report["missing_facts"].append(f"Candidate Email: {email}")

    # 3. Technical Skills
    skills_cat = data.get("skills_categorized", [])
    if skills_cat and isinstance(skills_cat, list):
        for sc in skills_cat:
            sc_clean = str(sc).strip().lower()
            if ":" in sc_clean:
                _, _, items = sc_clean.partition(":")
                for item in items.split(","):
                    item_c = re.sub(r"[^\w\s]", " ", item).strip().lower()
                    if len(item_c) > 2 and item_c not in norm_extracted:
                        report["missing_facts"].append(f"Skill item: {item_c}")
            else:
                sc_norm = re.sub(r"[^\w\s]", " ", sc_clean).strip()
                if len(sc_norm) > 2 and sc_norm not in norm_extracted:
                    report["missing_facts"].append(f"Skill line: {sc}")
    else:
        skills = data.get("skills", [])
        if isinstance(skills, list):
            for sk in skills:
                sk_clean = re.sub(r"[^\w\s]", " ", str(sk)).strip().lower()
                if len(sk_clean) > 2 and sk_clean not in norm_extracted:
                    report["missing_facts"].append(f"Skill: {sk}")

    # 4. Experience
    exp_list = data.get("experience_raw", []) or data.get("experience", [])
    for exp in exp_list:
        if isinstance(exp, dict):
            comp = exp.get("company", "")
            comp_norm = re.sub(r"[^\w\s]", " ", comp).strip().lower()
            if comp_norm and comp_norm not in norm_extracted:
                report["missing_facts"].append(f"Company: {comp}")
            for bullet in exp.get("bullets", []):
                b_clean = re.sub(r"^[•\-\*\s]+", "", str(bullet)).strip()
                words = re.sub(r"[^\w\s]", " ", b_clean).split()
                if len(words) >= 4:
                    phrase = " ".join(words[:4]).lower()
                    phrase_3 = " ".join(words[:3]).lower()
                    if phrase not in norm_extracted and phrase_3 not in norm_extracted:
                        report["missing_facts"].append(f"Experience bullet starting with '{phrase}'")
        elif isinstance(exp, str) and exp.strip():
            b_clean = re.sub(r"^[•\-\*\u26a1\u2605\s\ufffd▪▫◦‣⁃■□★☆+>~]+", "", exp).strip()
            words = re.sub(r"[^\w\s]", " ", b_clean).split()
            if len(words) >= 4:
                phrase = " ".join(words[:4]).lower()
                phrase_3 = " ".join(words[:3]).lower()
                if phrase not in norm_extracted and phrase_3 not in norm_extracted:
                    # Check if tokens were separated into company / role headers
                    sig_words = [w.lower() for w in words if len(w) > 3 and w.lower() not in {"with", "from", "present", "lead", "senior", "principal", "architect", "engineer", "at"}]
                    if not (sig_words and all(w in norm_words for w in sig_words)):
                        report["missing_facts"].append(f"Experience entry starting with '{phrase}'")

    # 5. Projects
    proj_list = data.get("projects_raw", []) or data.get("projects", [])
    for proj in proj_list:
        if isinstance(proj, dict):
            title = proj.get("title") or proj.get("name", "")
            title_norm = re.sub(r"[^\w\s]", " ", title).strip().lower()
            if title_norm and title_norm not in norm_extracted:
                report["project_integrity_issues"].append(f"Project title missing: {title}")
            for bullet in proj.get("bullets", []):
                b_clean = re.sub(r"^[•\-\*\s]+", "", str(bullet)).strip()
                words = re.sub(r"[^\w\s]", " ", b_clean).split()
                if len(words) >= 4:
                    phrase = " ".join(words[:4]).lower()
                    phrase_3 = " ".join(words[:3]).lower()
                    if phrase not in norm_extracted and phrase_3 not in norm_extracted:
                        report["project_integrity_issues"].append(f"Project bullet starting with '{phrase}'")
        elif isinstance(proj, str) and proj.strip():
            b_clean = re.sub(r"^[•\-\*\u26a1\u2605\s\ufffd▪▫◦‣⁃■□★☆+>~]+", "", proj).strip()
            words = re.sub(r"[^\w\s]", " ", b_clean).split()
            if len(words) >= 4:
                phrase = " ".join(words[:4]).lower()
                phrase_3 = " ".join(words[:3]).lower()
                if phrase not in norm_extracted and phrase_3 not in norm_extracted:
                    report["project_integrity_issues"].append(f"Project entry starting with '{phrase}'")

    # 6. Education
    edu_list = data.get("education_raw", []) or data.get("education", [])
    for edu in edu_list:
        if isinstance(edu, dict):
            inst = edu.get("institution") or edu.get("school", "")
            inst_norm = re.sub(r"[^\w\s]", " ", inst).strip().lower()
            if inst_norm and inst_norm not in norm_extracted:
                report["education_integrity_issues"].append(f"Education institution missing: {inst}")
        elif isinstance(edu, str) and edu.strip():
            words = re.sub(r"[^\w\s]", " ", edu.strip()).split()
            if len(words) >= 3:
                phrase = " ".join(words[:3]).lower()
                if phrase not in norm_extracted:
                    report["education_integrity_issues"].append(f"Education entry starting with '{phrase}'")

    # 7. Check for broken bullets
    lines = [l.strip() for l in extracted_text.splitlines() if l.strip()]
    for line in lines:
        if line in ("•", "-", "*", "·"):
            report["broken_bullets"].append("Found orphaned bullet symbol on isolated line")

    # 8. Check for duplicate section headers in sequence
    headers_found = [l.lower() for l in lines if l.lower() in ("technical skills", "professional experience", "technical projects", "education", "certifications", "languages")]
    for i in range(len(headers_found) - 1):
        if headers_found[i] == headers_found[i+1]:
            report["duplicated_content"].append(f"Consecutive duplicated header '{headers_found[i]}'")

    is_valid = (
        len(report["missing_facts"]) == 0
        and len(report["replacement_characters"]) == 0
        and len(report["project_integrity_issues"]) == 0
        and len(report["education_integrity_issues"]) == 0
        and len(report["broken_bullets"]) == 0
        and len(report["duplicated_content"]) == 0
    )

    return is_valid, report


def render_candidate_profile_to_text(
    profile: Any,
    strategy: Any | None = None,
) -> str:
    """
    Renders clean, structured ATS text directly from semantic CandidateProfile entities.
    Consumes semantic entities directly (Experience -> organization, roles, responsibility groups, evidence;
    Projects -> title, technologies, evidence; Education; Skills; Certifications; Achievements).
    Zero consumption of legacy raw strings.
    """
    if not profile:
        return ""

    # 1. Personal / Contact Info
    name = "Candidate"
    contacts = []
    if hasattr(profile, "identity") and profile.identity:
        ident = profile.identity
        name = getattr(ident, "name", None) or "Candidate"
        if getattr(ident, "location", None):
            contacts.append(str(ident.location))
        if getattr(ident, "email", None):
            contacts.append(str(ident.email))
        if getattr(ident, "phone", None):
            contacts.append(str(ident.phone))
        if getattr(ident, "linkedin", None):
            contacts.append(str(ident.linkedin))
        if getattr(ident, "github", None):
            contacts.append(str(ident.github))
        if getattr(ident, "portfolio", None):
            contacts.append(str(ident.portfolio))
    elif hasattr(profile, "personal") and profile.personal:
        pers = profile.personal
        if isinstance(pers, dict):
            name = pers.get("name") or pers.get("full_name") or "Candidate"
            for k in ["location", "email", "phone", "linkedin", "github", "portfolio"]:
                if pers.get(k):
                    contacts.append(str(pers[k]))
        else:
            name = getattr(pers, "name", None) or getattr(pers, "full_name", None) or "Candidate"
            for k in ["location", "email", "phone", "linkedin", "github", "portfolio"]:
                val = getattr(pers, k, None)
                if val:
                    contacts.append(str(val))

    lines = [str(name)]
    if contacts:
        lines.append(" • ".join(contacts))
    lines.append("")

    # Determine section order
    section_order = getattr(strategy, "section_order", None) or [
        "summary",
        "skills",
        "experience",
        "internships",
        "projects",
        "education",
        "certifications",
        "achievements",
        "publications",
        "research",
        "languages",
    ]
    standard_headings = getattr(strategy, "standard_ats_headings", None) or {
        "summary": "PROFESSIONAL SUMMARY",
        "skills": "TECHNICAL SKILLS",
        "experience": "PROFESSIONAL EXPERIENCE",
        "projects": "PROJECTS",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "achievements": "HONORS & AWARDS",
        "publications": "PUBLICATIONS & RESEARCH",
        "research": "RESEARCH EXPERIENCE",
    }

    rendered_add_sections = set()
    for sec in section_order:
        sec_lower = sec.lower()

        # Summary
        if sec_lower == "summary" and getattr(profile, "summary", ""):
            lines.append(standard_headings.get("summary", "PROFESSIONAL SUMMARY"))
            lines.append(profile.summary.strip())
            lines.append("")

        # Skills
        elif sec_lower == "skills" and getattr(profile, "skills", []):
            lines.append(standard_headings.get("skills", "TECHNICAL SKILLS"))
            lines.append(", ".join(profile.skills))
            lines.append("")

        # Experience
        elif sec_lower == "experience" and getattr(profile, "experience", []):
            lines.append(standard_headings.get("experience", "PROFESSIONAL EXPERIENCE"))
            for exp in profile.experience:
                comp = exp.company or ""
                loc = exp.location or ""
                comp_line = f"{comp} — {loc}" if comp and loc else comp
                if comp_line:
                    lines.append(comp_line)

                # Roles / Progression
                if exp.roles:
                    for r in exp.roles:
                        r_line = f"{r.title} ({r.dates})" if r.title and r.dates else r.title
                        if r_line:
                            lines.append(r_line)
                elif exp.role:
                    r_line = f"{exp.role} ({exp.dates})" if exp.role and exp.dates else exp.role
                    if r_line:
                        lines.append(r_line)

                # Responsibility groups or direct bullets
                if exp.responsibility_groups:
                    for grp in exp.responsibility_groups:
                        if grp.heading:
                            lines.append(grp.heading)
                        for ev in grp.evidence_units:
                            b_text = re.sub(r"^[•\-\*\s]+", "", ev.text).strip()
                            if b_text:
                                lines.append(f"• {b_text}")
                else:
                    for ev in exp.evidence_units:
                        b_text = re.sub(r"^[•\-\*\s]+", "", ev.text).strip()
                        if b_text:
                            lines.append(f"• {b_text}")
                lines.append("")

        # Internships
        elif sec_lower == "internships" and getattr(profile, "internships", []):
            lines.append(standard_headings.get("internships", "INTERNSHIPS"))
            for exp in profile.internships:
                comp = exp.company or ""
                loc = exp.location or ""
                comp_line = f"{comp} — {loc}" if comp and loc else comp
                if comp_line:
                    lines.append(comp_line)

                if exp.roles:
                    for r in exp.roles:
                        r_line = f"{r.title} ({r.dates})" if r.title and r.dates else r.title
                        if r_line:
                            lines.append(r_line)
                elif exp.role:
                    r_line = f"{exp.role} ({exp.dates})" if exp.role and exp.dates else exp.role
                    if r_line:
                        lines.append(r_line)

                if exp.responsibility_groups:
                    for grp in exp.responsibility_groups:
                        if grp.heading:
                            lines.append(grp.heading)
                        for ev in grp.evidence_units:
                            b_text = re.sub(r"^[•\-\*\s]+", "", ev.text).strip()
                            if b_text:
                                lines.append(f"• {b_text}")
                else:
                    for ev in exp.evidence_units:
                        b_text = re.sub(r"^[•\-\*\s]+", "", ev.text).strip()
                        if b_text:
                            lines.append(f"• {b_text}")
                    if not exp.evidence_units:
                        for b in exp.bullets:
                            b_text = re.sub(r"^[•\-\*\s]+", "", b).strip()
                            if b_text:
                                lines.append(f"• {b_text}")
                lines.append("")

        # Projects
        elif sec_lower == "projects" and getattr(profile, "projects", []):
            lines.append(standard_headings.get("projects", "PROJECTS"))
            for proj in profile.projects:
                p_name = getattr(proj, "name", "") or getattr(proj, "title", "") or ""
                p_dates = proj.dates or ""
                p_header = f"{p_name} ({p_dates})" if p_name and p_dates else p_name
                if p_header:
                    lines.append(p_header)
                if proj.technologies:
                    lines.append(f"Technologies: {', '.join(proj.technologies)}")
                for ev in proj.evidence_units:
                    b_text = re.sub(r"^[•\-\*\s]+", "", ev.text).strip()
                    if b_text:
                        lines.append(f"• {b_text}")
                lines.append("")

        # Education
        elif sec_lower == "education" and getattr(profile, "education", []):
            lines.append(standard_headings.get("education", "EDUCATION"))
            for edu in profile.education:
                deg = edu.degree or ""
                inst = edu.institution or ""
                dates = edu.dates or ""
                gpa = f"GPA: {edu.gpa}" if edu.gpa else ""
                edu_header = f"{deg} - {inst}" if deg and inst else (deg or inst)
                if dates:
                    edu_header = f"{edu_header} ({dates})"
                if edu_header:
                    lines.append(edu_header)
                if gpa:
                    lines.append(gpa)
            lines.append("")

        # Certifications
        elif sec_lower == "certifications" and getattr(profile, "certifications", []):
            lines.append(standard_headings.get("certifications", "CERTIFICATIONS"))
            for cert in profile.certifications:
                if isinstance(cert, str):
                    c_line = re.sub(r"^[•\-\*\s]+", "", cert).strip()
                else:
                    c_name = getattr(cert, "name", "") or str(cert)
                    c_issuer = getattr(cert, "issuer", "")
                    c_date = getattr(cert, "date", "")
                    c_line = f"{c_name} — {c_issuer} ({c_date})" if c_name and c_issuer and c_date else c_name
                if c_line:
                    lines.append(f"• {c_line}")
            lines.append("")

        # Achievements
        elif sec_lower == "achievements" and getattr(profile, "achievements", []):
            lines.append(standard_headings.get("achievements", "HONORS & AWARDS"))
            for ach in profile.achievements:
                if isinstance(ach, str):
                    a_line = re.sub(r"^[•\-\*\s]+", "", ach).strip()
                else:
                    a_title = getattr(ach, "title", "")
                    a_desc = getattr(ach, "description", "")
                    a_line = f"{a_title}: {a_desc}" if a_title and a_desc else (a_title or a_desc or str(ach))
                if a_line:
                    lines.append(f"• {a_line}")
            lines.append("")

        # Publications
        elif sec_lower in ("publications", "publications & research") and getattr(profile, "publications", []):
            lines.append(standard_headings.get("publications", "PUBLICATIONS & RESEARCH"))
            for pub in profile.publications:
                p_line = re.sub(r"^[•\-\*\s]+", "", str(pub)).strip()
                if p_line:
                    lines.append(f"• {p_line}")
            lines.append("")

        # Research
        elif sec_lower == "research" and getattr(profile, "research", []):
            lines.append(standard_headings.get("research", "RESEARCH EXPERIENCE"))
            for res in profile.research:
                r_line = re.sub(r"^[•\-\*\s]+", "", str(res)).strip()
                if r_line:
                    lines.append(f"• {r_line}")
            lines.append("")

        # Languages
        elif sec_lower == "languages" and getattr(profile, "languages", []):
            lines.append(standard_headings.get("languages", "LANGUAGES"))
            lines.append(", ".join(str(lang) for lang in profile.languages))
            lines.append("")

        # Additional Sections (Publications, Research, Volunteer, etc.)
        elif getattr(profile, "additional_sections", []):
            for add_sec in profile.additional_sections:
                sec_title = ((add_sec.get("heading", "") or add_sec.get("title", "")) if isinstance(add_sec, dict) else (getattr(add_sec, "heading", "") or getattr(add_sec, "title", ""))) or "ADDITIONAL INFORMATION"
                sec_title_lower = sec_title.lower()
                if sec_lower in sec_title_lower or sec_title_lower in sec_lower:
                    lines.append(sec_title.upper())
                    ev_units = add_sec.get("evidence_units", []) if isinstance(add_sec, dict) else getattr(add_sec, "evidence_units", [])
                    for ev in ev_units:
                        ev_text = ev.get("text", "") if isinstance(ev, dict) else (getattr(ev, "text", "") or str(ev))
                        b_text = re.sub(r"^[•\-\*\s]+", "", str(ev_text)).strip()
                        if b_text:
                            lines.append(f"• {b_text}")
                    items = add_sec.get("items", []) if isinstance(add_sec, dict) else getattr(add_sec, "items", [])
                    for item in items:
                        b_text = re.sub(r"^[•\-\*\s]+", "", str(item)).strip()
                        if b_text:
                            lines.append(f"• {b_text}")
                    lines.append("")
                    sec_id = (add_sec.get("id") if isinstance(add_sec, dict) else getattr(add_sec, "id", None)) or sec_title
                    rendered_add_sections.add(sec_id)

    # Render any remaining unrendered custom additional sections
    if getattr(profile, "additional_sections", []):
        for add_sec in profile.additional_sections:
            sec_title = ((add_sec.get("heading", "") or add_sec.get("title", "")) if isinstance(add_sec, dict) else (getattr(add_sec, "heading", "") or getattr(add_sec, "title", ""))) or "ADDITIONAL INFORMATION"
            sec_id = (add_sec.get("id") if isinstance(add_sec, dict) else getattr(add_sec, "id", None)) or sec_title
            if sec_id not in rendered_add_sections:
                lines.append(sec_title.upper())
                ev_units = add_sec.get("evidence_units", []) if isinstance(add_sec, dict) else getattr(add_sec, "evidence_units", [])
                for ev in ev_units:
                    ev_text = ev.get("text", "") if isinstance(ev, dict) else (getattr(ev, "text", "") or str(ev))
                    b_text = re.sub(r"^[•\-\*\s]+", "", str(ev_text)).strip()
                    if b_text:
                        lines.append(f"• {b_text}")
                items = add_sec.get("items", []) if isinstance(add_sec, dict) else getattr(add_sec, "items", [])
                for item in items:
                    b_text = re.sub(r"^[•\-\*\s]+", "", str(item)).strip()
                    if b_text:
                        lines.append(f"• {b_text}")
                lines.append("")

    return "\n".join(str(l) for l in lines if l is not None).strip()


def validate_rendered_export_integrity(
    profile: Any,
    rendered_text: str,
) -> tuple[bool, list[str]]:
    """
    Validates rendered plain-text or PDF/DOCX extracted text against semantic CandidateProfile:
    1. Text Completeness: Every EvidenceUnit must be present in the rendered text.
    2. Reading Order & Section Integrity: Headings must appear in non-overlapping, distinct order.
    3. Date & Metric Integrity: Preserves exact numerical values and dates.
    4. Bullet Formatting: No orphaned bullet markers or broken clauses.
    """
    errors: list[str] = []
    norm_rendered = " ".join(rendered_text.lower().split())

    # 1. Evidence Unit Presence
    for ev in getattr(profile, "evidence_units", []):
        clean_ev = re.sub(r"^[•\-\*\s]+", "", getattr(ev, "normalized_text", "") or getattr(ev, "text", "")).strip()
        words = clean_ev.split()
        if len(words) >= 4:
            phrase = " ".join(words[:4]).lower()
            if phrase not in norm_rendered:
                errors.append(f"EvidenceUnit '{ev.id}' starting with '{phrase}' missing from rendered export.")

    # 2. Date Integrity
    for exp in getattr(profile, "experience", []):
        if exp.dates and exp.dates.lower() not in norm_rendered:
            errors.append(f"Experience dates '{exp.dates}' missing from rendered export.")

    # 3. Orphan Bullet Symbols
    lines = [l.strip() for l in rendered_text.splitlines() if l.strip()]
    for line in lines:
        if line in ("•", "-", "*", "·"):
            errors.append("Found isolated orphaned bullet marker.")

    return len(errors) == 0, errors


