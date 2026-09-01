"""
Raw text (+ layout) extraction from uploaded resume files.
Deterministic, format-independent document extraction for PDF and DOCX.
Extracts font styles, layout geometry, reading order, and creates a canonical NormalizedDocument.
"""
import io
import fitz  # PyMuPDF
from docx import Document
from typing import Any

from app.modules.resume.parsing.document import (
    DocumentSpan,
    DocumentLine,
    DocumentBlock,
    DocumentPage,
    NormalizedDocument,
)
from app.modules.resume.parsing.normalization import (
    normalize_unicode_artifacts,
    detect_and_order_columns,
    suppress_repeated_headers_footers,
)


class UnsupportedFileTypeError(Exception):
    pass


class CorruptedFileError(Exception):
    pass


def extract_pdf_to_document(file_bytes: bytes) -> tuple[str, list[dict[str, Any]], NormalizedDocument, bool]:
    """
    Extracts text, layout spans, and geometry from PDF bytes into a canonical NormalizedDocument.
    Returns: (full_text, legacy_blocks, normalized_document, is_scanned)
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise CorruptedFileError(f"Could not open PDF: {exc}") from exc

    pages: list[DocumentPage] = []
    all_blocks: list[DocumentBlock] = []
    total_char_count = 0
    total_image_count = 0

    for page_num, page in enumerate(doc):
        rect = page.rect
        page_width = float(rect.width)
        page_height = float(rect.height)
        total_image_count += len(page.get_images())

        # Extract structured dictionary with blocks, lines, and spans
        page_dict = page.get_text("dict", sort=True)
        raw_blocks = page_dict.get("blocks", [])

        page_doc_blocks: list[DocumentBlock] = []

        for b_idx, b in enumerate(raw_blocks):
            if b.get("type") == 0:  # Text block
                bbox = tuple(float(x) for x in b.get("bbox", (0, 0, 0, 0)))
                doc_lines: list[DocumentLine] = []
                block_text_lines: list[str] = []

                for l in b.get("lines", []):
                    l_bbox = tuple(float(x) for x in l.get("bbox", (0, 0, 0, 0)))
                    spans: list[DocumentSpan] = []
                    line_text_parts: list[str] = []

                    for s in l.get("spans", []):
                        s_text = s.get("text", "")
                        if s_text:
                            line_text_parts.append(s_text)
                            total_char_count += len(s_text.strip())
                            flags = s.get("flags", 0)
                            is_bold = bool(flags & 2 or flags & 16 or "bold" in s.get("font", "").lower())
                            is_italic = bool(flags & 1 or "italic" in s.get("font", "").lower() or "oblique" in s.get("font", "").lower())
                            spans.append(DocumentSpan(
                                text=s_text,
                                font=s.get("font"),
                                size=float(s.get("size", 0.0)),
                                is_bold=is_bold,
                                is_italic=is_italic,
                                bbox=tuple(float(x) for x in s.get("bbox", (0, 0, 0, 0))),
                            ))

                    line_raw_text = "".join(line_text_parts).strip()
                    if line_raw_text:
                        line_norm = normalize_unicode_artifacts(line_raw_text)
                        has_bullet = line_norm.startswith("•") or line_raw_text.startswith(("-", "*", "•"))
                        doc_lines.append(DocumentLine(
                            text=line_raw_text,
                            normalized_text=line_norm,
                            spans=spans,
                            bbox=l_bbox,
                            has_bullet=has_bullet,
                        ))
                        block_text_lines.append(line_norm)

                raw_block_text = "\n".join(block_text_lines).strip()
                if raw_block_text:
                    norm_block_text = normalize_unicode_artifacts(raw_block_text)
                    doc_block = DocumentBlock(
                        id=f"p{page_num}_b{b_idx}",
                        page=page_num,
                        text=raw_block_text,
                        normalized_text=norm_block_text,
                        lines=doc_lines,
                        bbox=bbox,
                    )
                    page_doc_blocks.append(doc_block)

        # Infer reading order and handle multi-column layout for this page
        ordered_blocks = detect_and_order_columns(page_doc_blocks, page_width=page_width)

        # Determine if page is multi-column
        col0_count = sum(1 for b in ordered_blocks if b.column_index == 0)
        col1_count = sum(1 for b in ordered_blocks if b.column_index == 1)
        is_multi_col = col0_count >= 2 and col1_count >= 2

        doc_page = DocumentPage(
            page_number=page_num,
            width=page_width,
            height=page_height,
            blocks=ordered_blocks,
            is_multi_column=is_multi_col,
            column_count=2 if is_multi_col else 1,
        )
        pages.append(doc_page)
        all_blocks.extend(ordered_blocks)

    doc.close()

    # Detect repeated headers / footers across multi-page resumes
    suppress_repeated_headers_footers(pages)

    # Scanned PDF detection heuristic: 0 text or fewer than 30 chars with images present
    is_scanned = total_char_count < 30 and total_image_count > 0

    # Build canonical normalized text from non-suppressed blocks in reading order
    text_parts = []
    for p in pages:
        for b in p.blocks:
            if not b.is_repeated_header_or_footer:
                t = b.normalized_text or b.text
                if t.strip():
                    text_parts.append(t.strip())

    full_text = "\n".join(text_parts)

    norm_doc = NormalizedDocument(
        pages=pages,
        full_text=full_text,
        normalized_text=full_text,
        is_scanned=is_scanned,
        has_tables=False,
        file_type="pdf",
        blocks=all_blocks,
    )

    legacy_blocks = norm_doc.to_legacy_blocks()
    return full_text, legacy_blocks, norm_doc, is_scanned


def extract_docx_to_document(file_bytes: bytes) -> tuple[str, list[dict[str, Any]], bool, NormalizedDocument]:
    """
    Extracts text, paragraphs, lists, and tables from DOCX into a canonical NormalizedDocument.
    Returns: (full_text, legacy_blocks, has_tables, normalized_document)
    """
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise CorruptedFileError(f"Could not open DOCX: {exc}") from exc

    doc_blocks: list[DocumentBlock] = []
    text_parts: list[str] = []
    has_tables = len(document.tables) > 0

    # 1. Extract Paragraphs
    for p_idx, p in enumerate(document.paragraphs):
        p_text = p.text.strip()
        if not p_text:
            continue

        spans: list[DocumentSpan] = []
        for r in p.runs:
            if r.text:
                spans.append(DocumentSpan(
                    text=r.text,
                    font=r.font.name if r.font else None,
                    size=float(r.font.size.pt) if (r.font and r.font.size) else None,
                    is_bold=bool(r.bold),
                    is_italic=bool(r.italic),
                ))

        norm_p_text = normalize_unicode_artifacts(p_text)
        has_bullet = norm_p_text.startswith("•") or p_text.startswith(("-", "*", "•")) or (p.style and "list" in p.style.name.lower())

        doc_line = DocumentLine(
            text=p_text,
            normalized_text=norm_p_text,
            spans=spans,
            bbox=(0.0, float(p_idx * 15), 500.0, float(p_idx * 15 + 14)),
            has_bullet=has_bullet,
        )

        doc_block = DocumentBlock(
            id=f"docx_p_{p_idx}",
            page=0,
            text=p_text,
            normalized_text=norm_p_text,
            lines=[doc_line],
            bbox=(0.0, float(p_idx * 15), 500.0, float(p_idx * 15 + 14)),
            reading_order_index=len(doc_blocks),
        )
        doc_blocks.append(doc_block)
        text_parts.append(norm_p_text)

    # 2. Extract Tables
    for t_idx, table in enumerate(document.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                c_text = cell.text.strip()
                if c_text:
                    norm_c = normalize_unicode_artifacts(c_text)
                    if norm_c not in text_parts:
                        b_id = f"docx_t{t_idx}_r{r_idx}_c{c_idx}"
                        doc_block = DocumentBlock(
                            id=b_id,
                            page=0,
                            text=c_text,
                            normalized_text=norm_c,
                            lines=[DocumentLine(text=c_text, normalized_text=norm_c)],
                            bbox=(0.0, float(len(doc_blocks) * 15), 500.0, float(len(doc_blocks) * 15 + 14)),
                            reading_order_index=len(doc_blocks),
                            is_table_cell=True,
                        )
                        doc_blocks.append(doc_block)
                        text_parts.append(norm_c)

    full_text = "\n".join(text_parts)

    page_0 = DocumentPage(
        page_number=0,
        width=612.0,
        height=792.0,
        blocks=doc_blocks,
        is_multi_column=False,
        column_count=1,
    )

    norm_doc = NormalizedDocument(
        pages=[page_0],
        full_text=full_text,
        normalized_text=full_text,
        is_scanned=False,
        has_tables=has_tables,
        file_type="docx",
        blocks=doc_blocks,
    )

    legacy_blocks = norm_doc.to_legacy_blocks()
    return full_text, legacy_blocks, has_tables, norm_doc


def extract_pdf(file_bytes: bytes) -> tuple[str, list[dict[str, Any]]]:
    """
    Backward-compatible 2-tuple wrapper for PDF extraction.
    Returns: (full_text, legacy_blocks)
    """
    text, blocks, _, _ = extract_pdf_to_document(file_bytes)
    return text, blocks


def extract_docx(file_bytes: bytes) -> tuple[str, list[dict[str, Any]], bool]:
    """
    Backward-compatible 3-tuple wrapper for DOCX extraction.
    Returns: (full_text, legacy_blocks, has_tables)
    """
    text, blocks, has_tables, _ = extract_docx_to_document(file_bytes)
    return text, blocks, has_tables


def extract_text_and_layout(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """
    Unified extraction entry point.
    Returns:
    {
        "text": str,
        "blocks": list[dict],
        "has_tables": bool | None,
        "file_type": "pdf" | "docx",
        "document": NormalizedDocument,
        "is_scanned": bool,
    }
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        if not file_bytes.startswith(b"%PDF"):
            raise CorruptedFileError("File does not have a valid PDF signature.")
        text, blocks, doc_obj, is_scanned = extract_pdf_to_document(file_bytes)
        return {
            "text": text,
            "blocks": blocks,
            "has_tables": False,
            "file_type": "pdf",
            "document": doc_obj,
            "is_scanned": is_scanned,
        }

    if lower.endswith(".docx"):
        if not file_bytes.startswith(b"PK"):
            raise CorruptedFileError("File does not have a valid DOCX (zip) signature.")
        text, blocks, has_tables, doc_obj = extract_docx_to_document(file_bytes)
        return {
            "text": text,
            "blocks": blocks,
            "has_tables": has_tables,
            "file_type": "docx",
            "document": doc_obj,
            "is_scanned": False,
        }

    raise UnsupportedFileTypeError(f"Unsupported file type for '{filename}'. Only PDF and DOCX are supported.")
